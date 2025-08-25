########################################################################
#   SYNOP Data Processing Script                                       #
#                                                                      #
#   This script is designed to process SYNOP (Surface Synoptic         #
#   Observations) data files in BUFR format. It extracts relevant      #
#   information, performs data manipulation, and exports the processed #
#   data to an Excel and Doc file.                                     #
#                                                                      #
#   AUTHORS:                                                           #
#   - Issam LAGHA                                                      #
#   - Nour El Isslam KERROUMI                                          #
#                                                                      #
#   LAST MODIFICATION:                                                 #
#   - Date: 24th March 2024                                            #
#   - Date: 02nd April 2024  - Creates Doc file, By Issam LAGHA        #
#                                                                      #
########################################################################

import sys
import os
import pdbufr
import pandas as pd
import numpy as np 
from openpyxl import load_workbook
from openpyxl.styles import Border, Side

# Get environment variables for the current and previous day
AA = os.environ.get('AA')
MM = os.environ.get('MM')
DD = os.environ.get('DD')
AAprec = os.environ.get('AAprec')
MMprec = os.environ.get('MMprec')
DDprec = os.environ.get('DDprec')
PWD = os.environ.get('PWD') 
bufr_18 = f'{PWD}/../synop_alg_{AAprec}{MMprec}{DDprec}1800.bufr'
bufr_06 = f'{PWD}/../synop_alg_{AA}{MM}{DD}0600.bufr'


df_06 = pdbufr.read_bufr(bufr_06, columns=("stationOrSiteName",
                                          "heightOfStationGroundAboveMeanSeaLevel",
                                          "windDirection",
                                          "windSpeed",
                                          "cloudCoverTotal",
                                          "maximumTemperatureAtHeightAndOverPeriodSpecified",
                                          "minimumTemperatureAtHeightAndOverPeriodSpecified"
                                          ))
df_06_ww = pdbufr.read_bufr(bufr_06, columns=("stationOrSiteName",
                                          "presentWeather",
                                          "pastWeather1",
                                           "pastWeather2"
                                          ))

df_18 = pdbufr.read_bufr(bufr_18, columns=("stationOrSiteName",
                                        "heightOfStationGroundAboveMeanSeaLevel",
                                        "maximumTemperatureAtHeightAndOverPeriodSpecified"
                                        ),filters={"timePeriod": 0})
df_RR = pdbufr.read_bufr(bufr_06, columns=("stationOrSiteName",
                                          "heightOfStationGroundAboveMeanSeaLevel",
                                          "totalPrecipitationOrTotalWaterEquivalent"
                                          ),filters={"timePeriod": -24})
df_06.rename(columns={"stationOrSiteName":"STATIONS",
                   "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE EN METRES",
                   "windDirection":"Dir","windSpeed":"Vit (m/s)",
                   "cloudCoverTotal":"Néb (1/8)",
                   "totalPrecipitationOrTotalWaterEquivalent":"Précip (mm)",
                   "maximumTemperatureAtHeightAndOverPeriodSpecified":"Max de la veille",
                   "minimumTemperatureAtHeightAndOverPeriodSpecified":"Min de la nuit"
                   }, inplace=True)
df_RR.rename(columns={"stationOrSiteName":"STATIONS",
                     "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE",
                     "totalPrecipitationOrTotalWaterEquivalent":"Précip (mm)"
                     },inplace = True)
df_18.rename(columns={"stationOrSiteName":"STATIONS",
                     "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE",
                     "maximumTemperatureAtHeightAndOverPeriodSpecified":"Max"
                     },inplace = True)


df_06["Min de la nuit"] = (df_06["Min de la nuit"] - 273.15).round(1)
df_06["Néb (1/8)"] = (df_06["Néb (1/8)"] * 8 / 100).round()
df_06["ww"] = df_06_ww["presentWeather"]
df_06["w1"] = df_06_ww["pastWeather1"]
df_06["w2"] = df_06_ww["pastWeather2"]
df_06 = pd.merge(df_06, df_RR, on="STATIONS", how="inner")
df_RR = df_RR.dropna(subset = ["ALTITUDE"])
df_RR = df_RR.sort_values(["STATIONS"])
df_06 = df_06.dropna(subset = ["ALTITUDE EN METRES"])
df_06 = df_06.sort_values(["STATIONS"])
df_18 = df_18.dropna(subset = ["ALTITUDE"])
df_18 = df_18.sort_values(["STATIONS"])
pd.set_option('display.max_rows', 100)
df_merged = pd.merge(df_06, df_18, on="STATIONS", suffixes=('_06', '_18'), how='left')
df_merged["Max de la veille"]=df_merged["Max"] - 273.15
df_merged = df_merged[["STATIONS","ALTITUDE EN METRES","Dir","Vit (m/s)","Néb (1/8)","ww","Précip (mm)","w1","w2","Max de la veille","Min de la nuit"]]
rose_wind = {
    'N': list(range(338, 361)) + list(range(0, 23)),
    'NE': list(range(23, 68)),
    'E': list(range(68, 113)),
    'SE': list(range(113, 158)),
    'S': list(range(158, 203)),
    'SW': list(range(203, 248)),
    'W': list(range(248, 293)),
    'NW': list(range(293, 338))
}
# Function to convert degrees to 8-point rose wind
def convert_to_wind_rose(deg):
    if pd.isna(deg):
        return None  # Return None for NaN values
    for direction, sector in rose_wind.items():
        if int(deg) in sector:
            return direction
    return None  # Return None if degree does not match any sector
# Apply conversion and conditions to a DataFrame column 'Dir'
df_merged["Dir"] = df_merged["Dir"].apply(convert_to_wind_rose)

df_merged.loc[df_merged["Vit (m/s)"] == 0, "Dir"] = 'Calme'
df_merged.loc[(df_merged["Vit (m/s)"] < 2) & (df_merged["Vit (m/s)"] > 0), "Dir"] = 'VRB'

#df_06.to_csv('output_06.csv', index=True)
df_merged["ww"] = df_merged["ww"].where(df_merged["ww"] < 100, None)
df_merged["w1"] = df_merged["w1"].where(df_merged["w1"] < 10, None)
df_merged["w2"] = df_merged["w2"].where(df_merged["w2"] < 10, None)
df_merged["Précip (mm)"] = df_merged["Précip (mm)"].where(df_merged["Précip (mm)"] > 0, 0)

# Load Excel template
template_path = f'{PWD}/../templates/template.xlsx'
wb = load_workbook(template_path)
ws = wb.active

# Insert DataFrame into Excel template starting from 4th row
start_row = 4
for row in df_merged.itertuples(index=False):
    ws.append(row)

# Define border style
border_style = Border(left=Side(style='thin'),
                      right=Side(style='thin'),
                      top=Side(style='thin'),
                      bottom=Side(style='thin'))

# Apply border to all cells
for row in ws.iter_rows():

    for cell in row:
        cell.border = border_style

# Save to Excel
wb.save(f'{PWD}/../output.xlsx')

from openpyxl import load_workbook
from openpyxl.styles import Alignment
# Load data from output1.xlsx and list_stations.xlsx
output_df = pd.read_excel(f'{PWD}/../output.xlsx')
list_stations_wb = load_workbook(f'{PWD}/../templates/liste_stations_ref.xlsx')
list_stations_ws = list_stations_wb.active

# Iterate over list_stations.xlsx and search for each station in output1.xlsx
for row_idx, row in enumerate(list_stations_ws.iter_rows(), start=1):
    station_name = row[0].value  # Assuming STATIONS is in the first column
    if station_name:  # Check if station name is not empty
        # Find the row index of the station in output1.xlsx
        station_row_index = output_df.index[output_df['STATIONS'] == station_name].tolist()
        if station_row_index:
            # Write the data of the station from output1.xlsx to list_stations.xlsx
            station_row_index = station_row_index[0]  # Assuming there's only one match
            station_data = output_df.iloc[station_row_index].tolist()
            for col_idx, cell in enumerate(row, start=1):
                cell.value = station_data[col_idx - 1]
                cell.alignment = Alignment(horizontal='center', vertical='center')

# Save the modified list_stations.xlsx
list_stations_wb.save(f'{PWD}/../Bulletin_{AA}{MM}{DD}0600.xlsx')
print('Création du fichier Excel terminée avec succès!')

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# Read the Word document
doc = Document(f"{PWD}/../templates/Bulletin.docx")

# Extract tables from page 2 and 3
tables = []
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.text.strip() == "STATIONS":
                    tables.append(table)

# Extract relevant columns from df_06
relevant_columns = ["Dir", "Vit (m/s)", "Néb (1/8)", "ww", "Précip (mm)", "w1", "w2", "Max de la veille", "Min de la nuit"]

# Merge data into the tables
for table in tables:
    # Set font size and font family for the entire table
    table.style.font.size = Pt(10)
    table.style.font.name = 'Sansation'

    for row_idx, row in enumerate(table.rows):
        if row_idx > 0:  # Skip header row
            station = row.cells[0].text
            altitude = row.cells[1].text
            if station in df_merged["STATIONS"].values:
                for col_idx, col_name in enumerate(relevant_columns):
                    if col_name in df_merged.columns and col_name not in ["STATIONS", "ALTITUDE EN METRES"]:
                        cell_value = df_merged.loc[df_merged["STATIONS"] == station, col_name].values
                        if len(cell_value) > 0:
                            value = cell_value[0]
                            # Format numeric columns with one digit after the comma
                            if col_name in ["Vit (m/s)", "Précip (mm)", "Max de la veille", "Min de la nuit"]:
                                value = "{:.1f}".format(value) if not pd.isnull(value) else ""
                            if col_name in ["Néb (1/8)", "ww", "w1", "w2"]:
                                value = round(value) if not pd.isnull(value) else ""
                            #if col_name == "Néb (1/8)":  # Check for specific column
                            #   value = int(value)  # Convert to integer
                            #else:
                            #    value = "{:.1f}".format(value) if value != 0 else "0"
                            value = "" if pd.isnull(value) else value
                            row.cells[col_idx + 2].paragraphs[0].alignment = 1
                            cell = row.cells[col_idx + 2]
                            # Set font size and font family for each cell
                            cell.paragraphs[0].style.font.size = Pt(10)
                            cell.paragraphs[0].style.font.name = 'Sansation'
                            cell.text = str(value)

print('Insértion du tableau terminée avec succès!')
# Insert images
# Page 4
png_image_path1 = f"{PWD}/../geopotential_and_temperature_{AA}{MM}{DD}0000.png"
#png_image_path1 = f"{PWD}/../geopotential_and_temperature_{AA}{MM}{DD}0000.png"
paragraph_after_page4 = "Situation générale en altitude à 500 HPA à 00h TU"  # Text after which the image should be inserted
for paragraph in doc.paragraphs:
    if paragraph_after_page4 in paragraph.text:
        run = paragraph.add_run()
        run.add_picture(png_image_path1, width=Inches(6))  # Adjust width as needed
        doc.paragraphs[-1].alignment = 1
        # Add a text box under the picture
        #shape = doc.add_shape(
        #    MSO_SHAPE.RECTANGLE,
        #    Inches(0.5), Inches(0.5), Inches(6), Inches(1.5)
        #)
        #shape.text = "Analyse de la Situation en altitude :"  # Replace with your desired text
        #shape.fill.solid()
        #shape.fill.fore_color.rgb = RGBColor(255, 255, 255)  # Set background color of the text box
        #shape.line.color.rgb = RGBColor(0, 0, 0)  # Set border color of the text box
        #shape.line.width = Pt(1)  # Set border width of the text box

        # Center align the text inside the text box
        #for paragraph in shape.text_frame.paragraphs:
        #    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        break  # Exit loop after inserting the image

# Page 5
png_image_path2 = f"{PWD}/../mslp_{AA}{MM}{DD}0600.png"
paragraph_after_page5 = "Situation générale en Surface à 06h TU"  # Text after which the image should be inserted
for paragraph in doc.paragraphs:
    if paragraph_after_page5 in paragraph.text:
        run = paragraph.add_run()
        run.add_picture(png_image_path2, width=Inches(6))  # Adjust width as needed
        doc.paragraphs[-1].alignment = 1
        # Add a text box under the picture
        #shape = doc.add_shape(
        #    MSO_SHAPE.RECTANGLE,
        #    Inches(0.5), Inches(0.5), Inches(6), Inches(1.5)
        #)
        #shape.text = "Analyse de la situation au niveau de la mer :"  # Replace with your desired text
        #shape.fill.solid()
        #shape.fill.fore_color.rgb = RGBColor(255, 255, 255)  # Set background color of the text box
        #shape.line.color.rgb = RGBColor(0, 0, 0)  # Set border color of the text box
        #shape.line.width = Pt(1)  # Set border width of the text box

        # Center align the text inside the text box
        #for paragraph in shape.text_frame.paragraphs:
        #    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        break  # Exit loop after inserting the image

# Save the modified document
doc.save(f'{PWD}/../BQRM_{AA}{MM}{DD}0600.docx')
print('Insértion des images terminée avec succès!')
print('Création du fichier Doc terminée avec succès!')


