import sys
import os
import logging
import pandas as pd
import numpy as np
from pathlib import Path
import openpyxl
from openpyxl.styles import Border, Side, Alignment
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Optional, Union, Any
import argparse # <--- ADD THIS LINE

# Import functions from pytaps
from pytaps.logging_utils import setup_logger
from pytaps.data_utils import read_bufr_to_dataframe, load_excel_workbook

# --- Command-line argument parsing ---
parser = argparse.ArgumentParser(description="BufrToXLS_ref.py: Processes BUFR data to generate Excel and Word reports.")
# IMPORTANT: This script MUST receive the shared log file path from the previous script
parser.add_argument('--shared-log-file', type=str, required=True,
                    help='Path to a shared log file for centralized logging.')
args = parser.parse_args()

# --- Logger Setup (using pytaps) ---
script_name = Path(__file__).stem # Gets the script name without .py extension
# Configure the logger using the shared log file path received from the command line
logger, current_log_file_path = setup_logger(
    script_name=script_name,
    log_level=logging.INFO,
    shared_log_file_path=args.shared_log_file # <--- THIS IS THE KEY CHANGE
)

logger.info("-------------------------------------------------------------------")
logger.info(f"Script {script_name}.py started.")
logger.info(f"Logging to shared file: {current_log_file_path}")
logger.info("-------------------------------------------------------------------")
# --- End Logger Setup ---


# Get environment variables for the current and previous day
logger.info("Retrieving environment variables...")
# Use a dictionary for cleaner iteration and validation of environment variables
env_vars_needed = {
    'AA': None, 'MM': None, 'DD': None,
    'AAprec': None, 'MMprec': None, 'DDprec': None,
    'PWD': None
}

for var_name in env_vars_needed.keys():
    env_vars_needed[var_name] = os.environ.get(var_name)

missing_env_vars = [k for k, v in env_vars_needed.items() if v is None]

if missing_env_vars:
    logger.critical(f"Missing required environment variables: {', '.join(missing_env_vars)}. Exiting.")
    sys.exit(1)
else:
    # Assign back to individual variables for clarity (optional, can use env_vars_needed dict directly)
    AA = env_vars_needed['AA']
    MM = env_vars_needed['MM']
    DD = env_vars_needed['DD']
    AAprec = env_vars_needed['AAprec']
    MMprec = env_vars_needed['MMprec']
    DDprec = env_vars_needed['DDprec']
    PWD = env_vars_needed['PWD']
    logger.info("All required environment variables retrieved successfully.")
    logger.debug(f"AA: {AA}, MM: {MM}, DD: {DD}, AAprec: {AAprec}, MMprec: {MMprec}, DDprec: {DDprec}, PWD: {PWD}")

# Construct BUFR file paths using Path objects for robustness
base_dir = Path(PWD).parent # Go up one level from PWD (local_directory to bqrm_output_root)
bufr_18_path = base_dir / f'synop_alg_{AAprec}{MMprec}{DDprec}1800.bufr'
bufr_06_path = base_dir / f'synop_alg_{AA}{MM}{DD}0600.bufr'

logger.info(f"Constructed BUFR 18h file path: {bufr_18_path}")
logger.info(f"Constructed BUFR 06h file path: {bufr_06_path}")

# --- Read BUFR data using pytaps.data_utils.read_bufr_to_dataframe ---
try:
    logger.info(f"Attempting to read 06h BUFR data from {bufr_06_path} (main columns).")
    df_06 = read_bufr_to_dataframe(
        bufr_06_path,
        columns=["stationOrSiteName",
                 "heightOfStationGroundAboveMeanSeaLevel",
                 "windDirection",
                 "windSpeed",
                 "cloudCoverTotal",
                 "maximumTemperatureAtHeightAndOverPeriodSpecified",
                 "minimumTemperatureAtHeightAndOverPeriodSpecified"],
        logger_instance=logger
    )
    if df_06 is None:
        logger.critical("Failed to load main 06h BUFR data. Exiting.")
        sys.exit(1)
    logger.info(f"Successfully read {len(df_06)} records for df_06.")

    logger.info(f"Attempting to read 06h BUFR data from {bufr_06_path} (present and past weather).")
    df_06_ww = read_bufr_to_dataframe(
        bufr_06_path,
        columns=["stationOrSiteName", "presentWeather", "pastWeather1", "pastWeather2"],
        logger_instance=logger
    )
    if df_06_ww is None:
        logger.critical("Failed to load 06h BUFR weather data. Exiting.")
        sys.exit(1)
    logger.info(f"Successfully read {len(df_06_ww)} records for df_06_ww.")

    logger.info(f"Attempting to read 18h BUFR data from {bufr_18_path}.")
    df_18 = read_bufr_to_dataframe(
        bufr_18_path,
        columns=["stationOrSiteName", "heightOfStationGroundAboveMeanSeaLevel", "maximumTemperatureAtHeightAndOverPeriodSpecified"],
        filters={"timePeriod": 0}, # Pass filter directly to pdbufr via pytaps function
        logger_instance=logger
    )
    if df_18 is None:
        logger.critical("Failed to load 18h BUFR data. Exiting.")
        sys.exit(1)
    logger.info(f"Successfully read {len(df_18)} records for df_18.")

    logger.info(f"Attempting to read 06h BUFR data from {bufr_06_path} (precipitation).")
    df_RR = read_bufr_to_dataframe(
        bufr_06_path,
        columns=["stationOrSiteName", "heightOfStationGroundAboveMeanSeaLevel", "totalPrecipitationOrTotalWaterEquivalent"],
        filters={"timePeriod": -24}, # Pass filter directly to pdbufr via pytaps function
        logger_instance=logger
    )
    if df_RR is None:
        logger.critical("Failed to load 06h BUFR precipitation data. Exiting.")
        sys.exit(1)
    logger.info(f"Successfully read {len(df_RR)} records for df_RR.")

except Exception as e: # Catch any remaining exceptions from pytaps functions
    logger.critical(f"An error occurred during BUFR file reading: {e}", exc_info=True)
    sys.exit(1)

# --- Rename columns ---
logger.info("Renaming DataFrame columns...")
try:
    df_06.rename(columns={
        "stationOrSiteName":"STATIONS",
        "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE EN METRES",
        "windDirection":"Dir",
        "windSpeed":"Vit (m/s)",
        "cloudCoverTotal":"Néb (1/8)",
        "maximumTemperatureAtHeightAndOverPeriodSpecified":"Max de la veille", # This will be overwritten later by df_18's Max
        "minimumTemperatureAtHeightAndOverPeriodSpecified":"Min de la nuit"
    }, inplace=True)
    df_RR.rename(columns={
        "stationOrSiteName":"STATIONS",
        "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE", # This ALTITUDE column from RR will be dropped after merge
        "totalPrecipitationOrTotalWaterEquivalent":"Précip (mm)"
    },inplace = True)
    df_18.rename(columns={
        "stationOrSiteName":"STATIONS",
        "heightOfStationGroundAboveMeanSeaLevel":"ALTITUDE", # This ALTITUDE column from 18 will be dropped after merge
        "maximumTemperatureAtHeightAndOverPeriodSpecified":"Max"
    },inplace = True)
    logger.info("Columns renamed successfully for df_06, df_RR, and df_18.")
except Exception as e:
    logger.critical(f"Error during column renaming: {e}", exc_info=True)
    sys.exit(1)

# --- Data Processing and Merging ---
logger.info("Starting data processing and merging operations...")
try:
    df_06["Min de la nuit"] = (df_06["Min de la nuit"] - 273.15).round(1)
    df_06["Néb (1/8)"] = (df_06["Néb (1/8)"] * 8 / 100).round()
    logger.info("Temperature (Min de la nuit) and cloud cover (Néb (1/8)) calculations applied to df_06.")

    # Merge df_06_ww into df_06
    # Ensure 'stationOrSiteName' is renamed to 'STATIONS' in df_06_ww before merge
    df_06_ww_renamed = df_06_ww[["stationOrSiteName", "presentWeather", "pastWeather1", "pastWeather2"]].rename(columns={"stationOrSiteName":"STATIONS"})
    df_06 = pd.merge(df_06, df_06_ww_renamed, on="STATIONS", how="left")
    df_06.rename(columns={"presentWeather":"ww", "pastWeather1":"w1", "pastWeather2":"w2"}, inplace=True)
    logger.info("Present and past weather codes (ww, w1, w2) added to df_06.")

    # Merge df_RR into df_06 (only precipitation and station, as ALTITUDE is redundant)
    df_06 = pd.merge(df_06, df_RR[["STATIONS", "Précip (mm)"]], on="STATIONS", how="left")
    logger.info(f"df_06 merged with df_RR (precipitation data). New df_06 shape: {df_06.shape}")

    # Clean and sort individual dataframes (as per original logic, though some might be redundant after merges)
    df_RR = df_RR.dropna(subset = ["ALTITUDE"]) # This df_RR is not used after the merge into df_06
    df_RR = df_RR.sort_values(["STATIONS"])
    logger.debug(f"df_RR cleaned (dropna on ALTITUDE, sort_values by STATIONS). Shape: {df_RR.shape} (Note: this df_RR is not used further).")

    df_06 = df_06.dropna(subset = ["ALTITUDE EN METRES"])
    df_06 = df_06.sort_values(["STATIONS"])
    logger.info(f"df_06 cleaned (dropna on ALTITUDE EN METRES, sort_values by STATIONS). Shape: {df_06.shape}")

    df_18 = df_18.dropna(subset = ["ALTITUDE"])
    df_18 = df_18.sort_values(["STATIONS"])
    logger.info(f"df_18 cleaned (dropna on ALTITUDE, sort_values by STATIONS). Shape: {df_18.shape}")

    pd.set_option('display.max_rows', 100)
    logger.debug("Pandas display option 'display.max_rows' set to 100.")

    # Final merge with df_18 (only Max and station)
    df_merged = pd.merge(df_06, df_18[["STATIONS", "Max"]], on="STATIONS", how='left')
    logger.info(f"df_06 merged with df_18 into df_merged. Shape: {df_merged.shape}")

    df_merged["Max de la veille"]=df_merged["Max"] - 273.15
    df_merged.drop(columns=["Max"], inplace=True) # Drop the temporary 'Max' column
    logger.info("Max de la veille calculated in df_merged (Kelvin to Celsius conversion) and temporary 'Max' column dropped.")

    # Select and reorder columns for final output
    final_columns = ["STATIONS","ALTITUDE EN METRES","Dir","Vit (m/s)","Néb (1/8)","ww","Précip (mm)","w1","w2","Max de la veille","Min de la nuit"]
    df_merged = df_merged[final_columns]
    logger.info("df_merged columns reordered to final output format.")

    # Define wind rose mapping
    rose_wind = {
        'N': list(range(338, 361)) + list(range(0, 23)),
        'NE': list(range(23, 68)),
        'E': list(range(68, 113)),
        'SE': list(range(113, 158)),
        'S': list(range(158, 203)),
        'SW': list(range(203, 248)),
        'W': list(range(248, 293)),
        'NW': list(range(293, 338))
    }
    # Function to convert degrees to 8-point rose wind
    def convert_to_wind_rose(deg: Union[float, int]) -> Optional[str]:
        if pd.isna(deg):
            return None
        try:
            int_deg = int(deg)
        except (ValueError, TypeError):
            logger.warning(f"Could not convert wind direction '{deg}' to integer for wind rose conversion. Returning None.")
            return None
        for direction, sector in rose_wind.items():
            if int_deg in sector:
                return direction
        return None
    df_merged["Dir"] = df_merged["Dir"].apply(convert_to_wind_rose)
    logger.info("Wind direction (Dir) converted to 8-point rose format.")

    df_merged.loc[df_merged["Vit (m/s)"] == 0, "Dir"] = 'Calme'
    df_merged.loc[(df_merged["Vit (m/s)"] < 2) & (df_merged["Vit (m/s)"] > 0), "Dir"] = 'VRB'
    logger.info("Wind direction adjusted for 'Calme' (0 m/s) and 'VRB' (0-2 m/s) conditions.")

    # Clean/filter weather codes and precipitation
    df_merged["ww"] = df_merged["ww"].where(df_merged["ww"] < 100, None)
    df_merged["w1"] = df_merged["w1"].where(df_merged["w1"] < 10, None)
    df_merged["w2"] = df_merged["w2"].where(df_merged["w2"] < 10, None)
    df_merged["Précip (mm)"] = df_merged["Précip (mm)"].where(df_merged["Précip (mm)"] > 0, 0)
    logger.info("Weather codes (ww, w1, w2) and precipitation values cleaned/filtered.")

except Exception as e:
    logger.critical(f"An error occurred during data processing and merging: {e}", exc_info=True)
    sys.exit(1)

# --- Excel Output (output.xlsx - temporary file) ---
try:
    # The original template_path for output.xlsx is no longer directly used for writing data,
    # as we'll create the file directly from df_merged.
    # template_path = base_dir / 'templates' / 'template.xlsx' # This line can be kept if template has other uses, but not for direct data write here.
    output_excel_path = base_dir / 'output.xlsx'

    logger.info(f"Saving df_merged to temporary Excel file: {output_excel_path}")
    # Use pandas to_excel to write the DataFrame directly, including headers.
    # index=False prevents pandas from writing the DataFrame index as a column.
    df_merged.to_excel(output_excel_path, index=False)
    logger.info("Temporary Excel file saved successfully with correct headers.")

    # Now, load the workbook that pandas just created to apply custom borders.
    # This requires openpyxl directly.
    wb = openpyxl.load_workbook(output_excel_path)
    ws = wb.active # Get the active worksheet (which will be the one pandas wrote to)

    logger.info("Applying border style to all cells in the worksheet.")
    border_style = Border(left=Side(style='thin'),
                          right=Side(style='thin'),
                          top=Side(style='thin'),
                          bottom=Side(style='thin'))
    # Iterate over all rows (including the header row written by pandas) to apply borders
    for row in ws.iter_rows(min_row=1): # Start from min_row=1 to include headers
        for cell in row:
            cell.border = border_style
    logger.info("Borders applied successfully to all cells.")

    logger.info(f"Resaving temporary Excel file with borders to: {output_excel_path}")
    wb.save(output_excel_path)
    logger.info("Temporary Excel file (with borders) resaved successfully.")

except Exception as e:
    logger.critical(f"An error occurred during the first Excel output phase (output.xlsx): {e}", exc_info=True)
    sys.exit(1)

# The rest of your script, especially the "Excel Output (Bulletin_...xlsx - final file)"
# section, should now work correctly because output_df will have the right column names.


# --- Excel Output (Bulletin_...xlsx - final file) ---
try:
    bulletin_excel_path = base_dir / f'Bulletin_{AA}{MM}{DD}0600.xlsx'
    list_stations_template_path = base_dir / 'templates' / 'liste_stations_ref.xlsx'

    logger.info(f"Reading data from {output_excel_path} for final Excel generation.")
    output_df = pd.read_excel(output_excel_path) # pd.read_excel is suitable for reading a DataFrame
    logger.info(f"Successfully read {len(output_df)} records from {output_excel_path}.")

    logger.info(f"Loading list of stations template from: {list_stations_template_path}")
    # Use pytaps.data_utils.load_excel_workbook
    list_stations_wb, list_stations_ws = load_excel_workbook(list_stations_template_path, logger_instance=logger)
    logger.info("List of stations template loaded successfully.")

    logger.info("Iterating through station list and populating data into the template.")
    stations_found_count = 0
    
    # Define the columns from output_df that correspond to the template columns.
    # Reorder output_df to ensure column order matches expected template order for direct list conversion.
    expected_output_df_cols = [
        "STATIONS", "ALTITUDE EN METRES", "Dir", "Vit (m/s)", "Néb (1/8)", "ww",
        "Précip (mm)", "w1", "w2", "Max de la veille", "Min de la nuit"
    ]
    output_df_reordered = output_df[expected_output_df_cols]

    # Assuming first row is header in template, start iterating from the second row (min_row=2)
    for row_idx, row_cells in enumerate(list_stations_ws.iter_rows(min_row=2), start=2):
        station_name_cell = row_cells[0] # First cell in the row
        station_name = station_name_cell.value
        
        if station_name:
            station_data_row = output_df_reordered[output_df_reordered['STATIONS'] == station_name]
            
            if not station_data_row.empty:
                station_data_list = station_data_row.iloc[0].tolist()

                for col_idx_in_template, cell in enumerate(row_cells):
                    if col_idx_in_template < len(station_data_list):
                        value_to_write = station_data_list[col_idx_in_template]
                        
                        # Apply formatting based on column type/name (indices correspond to `expected_output_df_cols`)
                        if col_idx_in_template in [3, 6, 9, 10]: # Vit, Precip, Max, Min (0-indexed)
                            cell.value = "{:.1f}".format(value_to_write) if not pd.isnull(value_to_write) else ""
                        elif col_idx_in_template in [4, 5, 7, 8]: # Neb, ww, w1, w2
                            cell.value = round(value_to_write) if not pd.isnull(value_to_write) else ""
                        else: # For STATIONS, ALTITUDE EN METRES, Dir
                            cell.value = value_to_write if not pd.isnull(value_to_write) else ""
                        
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    else:
                        logger.debug(f"Skipping writing to cell {cell.coordinate} for station {station_name} as no corresponding data in output_df.")
                stations_found_count += 1
            else:
                logger.warning(f"Station '{station_name}' from list_stations_ref.xlsx not found in output data. Skipping.")
    logger.info(f"Finished populating station data. {stations_found_count} stations found and updated in the Excel template.")

    logger.info(f"Saving final Excel file to: {bulletin_excel_path}")
    list_stations_wb.save(bulletin_excel_path)
    logger.info('Création du fichier Excel terminée avec succès!')

except Exception as e:
    logger.critical(f"An unexpected error occurred during the second Excel output phase (Bulletin_...xlsx): {e}", exc_info=True)
    sys.exit(1)

# --- Word Document Generation ---
# Helper function to insert image after a specific paragraph
def insert_image_after_paragraph(doc: Document, image_path: Path, anchor_text: str, width_inches: float, logger_instance: logging.Logger) -> bool:
    """
    Inserts an image into a Word document after a paragraph containing specific anchor text.
    The image is placed in a new, centered paragraph at the end of the document.

    Args:
        doc (Document): The python-docx Document object.
        image_path (Path): Path to the image file.
        anchor_text (str): The text to search for to determine insertion point.
        width_inches (float): Width of the image in inches.
        logger_instance (logging.Logger): Logger instance for messages.

    Returns:
        bool: True if image was inserted, False otherwise.
    """
    inserted = False
    if not image_path.exists() or not image_path.is_file():
        logger_instance.warning(f"Image file not found: {image_path}. Skipping insertion.")
        return False

    # Find the anchor paragraph
    anchor_paragraph = None
    for paragraph in doc.paragraphs:
        if anchor_text in paragraph.text:
            anchor_paragraph = paragraph
            break

    if anchor_paragraph:
        # Create a new paragraph and insert the image into it.
        # This new paragraph will be added to the end of the document.
        # For precise insertion after `anchor_paragraph`, more advanced python-docx
        # manipulation (using `_p.insert_paragraph_before()`) would be needed,
        # which is beyond the scope of simple refactoring.
        # This approach matches the original script's effect of adding images at the "end".
        new_image_paragraph = doc.add_paragraph()
        new_image_run = new_image_paragraph.add_run()
        try:
            new_image_run.add_picture(str(image_path), width=Inches(width_inches))
            new_image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logger_instance.info(f"Image ({image_path.name}) inserted successfully. Placed at end of document.")
            inserted = True
        except Exception as e:
            logger_instance.error(f"Error inserting image {image_path}: {e}", exc_info=True)
    else:
        logger_instance.warning(f"Anchor paragraph '{anchor_text}' for image '{image_path.name}' not found.")
    
    return inserted


try:
    word_template_path = base_dir / "templates" / "Bulletin.docx"
    output_word_path = base_dir / f'BQRM_{AA}{MM}{DD}0600.docx'

    logger.info(f"Reading Word document template from: {word_template_path}")
    doc = Document(word_template_path)
    logger.info("Word document template loaded successfully.")

    logger.info("Extracting tables from the Word document for data insertion.")
    tables_to_update = []
    # Heuristic to find tables: look for "STATIONS" in the first few rows of each table.
    for table in doc.tables:
        found_station_header = False
        for row_idx, row in enumerate(table.rows):
            if row_idx >= 3: # Only check the first 3 rows for header
                break
            for cell in row.cells:
                if "STATIONS" in cell.text.strip().upper():
                    tables_to_update.append(table)
                    found_station_header = True
                    break
            if found_station_header:
                break
    logger.info(f"Found {len(tables_to_update)} tables to update in the Word document.")

    relevant_columns = ["Dir", "Vit (m/s)", "Néb (1/8)", "ww", "Précip (mm)", "w1", "w2", "Max de la veille", "Min de la nuit"]
    logger.info("Merging data from DataFrame into Word document tables.")
    for table_idx, table in enumerate(tables_to_update):
        logger.debug(f"Processing table {table_idx + 1}/{len(tables_to_update)} for data insertion.")
        
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0: # Assuming first row is header, skip it.
                continue
            
            # Assuming station name is in the first column of the table
            station_cell_text = row.cells[0].text.strip()
            
            if station_cell_text in df_merged["STATIONS"].values:
                station_data_row = df_merged.loc[df_merged["STATIONS"] == station_cell_text].iloc[0]

                for col_name_idx, col_name in enumerate(relevant_columns):
                    if col_name in station_data_row.index: # Check if column exists in the data row
                        value = station_data_row[col_name]
                        
                        # Format numeric columns
                        if col_name in ["Vit (m/s)", "Précip (mm)", "Max de la veille", "Min de la nuit"]:
                            value_str = "{:.1f}".format(value) if not pd.isnull(value) else ""
                        elif col_name in ["Néb (1/8)", "ww", "w1", "w2"]:
                            value_str = str(round(value)) if not pd.isnull(value) else ""
                        else:
                            value_str = str(value) if not pd.isnull(value) else "" # Handle non-numeric and NaN

                        # Column mapping: STATIONS (0), ALTITUDE EN METRES (1), then relevant_columns start at index 2
                        # So, table cell index is col_name_idx + 2
                        target_cell = row.cells[col_name_idx + 2] 
                        
                        # Clear existing content and add new run
                        if not target_cell.paragraphs:
                            target_cell.add_paragraph()
                        
                        cell_paragraph = target_cell.paragraphs[0]
                        # Clear all existing runs in the paragraph to replace content
                        for i in range(len(cell_paragraph.runs)):
                            cell_paragraph.runs[0].clear()

                        new_run = cell_paragraph.add_run(value_str)
                        new_run.font.size = Pt(10)
                        new_run.font.name = 'Sansation'
                        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                logger.debug(f"Station '{station_cell_text}' from Word document not found in merged DataFrame. Skipping row.")
    logger.info('Insértion du tableau terminée avec succès!')

    # Insert images
    logger.info("Inserting images into Word document.")
    
    png_image_path1 = base_dir / f"geopotential_and_temperature_{AA}{MM}{DD}0000.png"
    paragraph_after_page4 = "Situation générale en altitude à 500 HPA à 00h TU"
    insert_image_after_paragraph(doc, png_image_path1, paragraph_after_page4, 6, logger)

    png_image_path2 = base_dir / f"mslp_{AA}{MM}{DD}0600.png"
    paragraph_after_page5 = "Situation générale en Surface à 06h TU"
    insert_image_after_paragraph(doc, png_image_path2, paragraph_after_page5, 6, logger)

    logger.info(f"Saving final Word document to: {output_word_path}")
    doc.save(output_word_path)
    logger.info('Insértion des images terminée avec succès!')
    logger.info('Création du fichier Doc terminée avec succès!')

except Exception as e:
    logger.critical(f"An unexpected error occurred during Word document generation: {e}", exc_info=True)
    sys.exit(1)

logger.info("Script finished successfully!")
