import json
import os
import pandas as pd
from datetime import datetime, timedelta
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from typing import Dict, List, Union, Any, Tuple
from pathlib import Path

# Import functions from your pytaps package
from pytaps.logging_utils import setup_logger
from pytaps.file_operations import check_file_exists_and_log, ensure_parent_directory_exists
from pytaps.data_utils import read_bufr_to_dataframe, load_excel_workbook, save_dataframe_to_csv # save_dataframe_to_csv not strictly needed but good to show

import argparse
import sys

parser = argparse.ArgumentParser(description=f"{os.path.basename(__file__)} script.")


# Add positional arguments specific to this script (still needed)
parser.add_argument('year', type=str, help='Year (YYYY) for data processing.')
parser.add_argument('month', type=str, help='Month (MM) for data processing.')
parser.add_argument('day', type=str, help='Day (DD) for data processing.')
parser.add_argument('local_directory', type=str, help='Local base directory for file operations.')

# --- Debugging Prints for Argument Parsing ---
try:
    # Now parse_args() should work directly as --shared-log-file is no longer expected
    args = parser.parse_args() 
except SystemExit as e:
    print(f"ERROR ({os.path.basename(__file__)}): Argument parsing failed. Exit code: {e.code}", file=sys.stderr)
    sys.exit(e.code)

# --- NEW BLOCK: Get shared log file path from environment variable ---
shared_log_file_from_env = os.environ.get('PYTAPS_SHARED_LOG_FILE')
if shared_log_file_from_env:
    logger_setup_message = f"Shared log file path from environment: {shared_log_file_from_env}"
    # print(f"DEBUG ({os.path.basename(__file__)}): {logger_setup_message}", file=sys.stderr) # Keep for initial debugging, remove later
else:
    logger_setup_message = "PYTAPS_SHARED_LOG_FILE environment variable not set. This script will create its own log file."
    # print(f"WARNING ({os.path.basename(__file__)}): {logger_setup_message}", file=sys.stderr) # Keep for initial debugging, remove later
# --- END NEW BLOCK ---

script_name_for_log = os.path.basename(__file__).replace(".py", "")
logger, actual_log_path_from_pytaps = setup_logger(
    script_name=script_name_for_log,
    log_directory_base=os.path.dirname(os.path.abspath(__file__)),
    log_level=logging.INFO,
    shared_log_file_path=shared_log_file_from_env # Use the path from the environment variable
)
logger.info(logger_setup_message) # Log the message about source of log path


logger.info("Script execution started.")

# --- 2. Environment Variables ---
# Keep this section as is, as it's core configuration, but use the new logger
AA = os.environ.get('AA')
MM = os.environ.get('MM')
DD = os.environ.get('DD')
PWD = os.environ.get('PWD')

if not all([AA, MM, DD, PWD]):
    logger.critical("One or more environment variables (AA, MM, DD, PWD) are not set. Exiting.")
    exit(1) # Use exit(1) for error exit code
logger.info(f"Environment variables loaded: AA={AA}, MM={MM}, DD={DD}, PWD={PWD}")

# --- 3. File Paths (using pathlib for better path handling) ---
# Convert PWD to a Path object for easier manipulation
base_dir = Path(PWD)

bufr_file = base_dir / f'Synop/Synop_{AA}{MM}{DD}0600.bufr'
xlsx_24h_template_path = base_dir / 'templates/cumul24.xlsx'
xlsx_agri_template_path = base_dir / 'templates/agricole.xlsx'
output_xlsx_24h_path = base_dir / f'cumul24_{MM}{DD}0600.xlsx'
output_xlsx_agri_path = base_dir / f'agricole_{MM}{DD}0600.xlsx'
word_template_path = base_dir / 'templates/cumul.docx'
output_word_path = base_dir / f'Cumul_table{MM}{DD}0600.docx'
mapping_file = base_dir / 'ListStation.json'
missing_stations_path = base_dir / f"Synop{AA}{MM}{DD}-18.txt"
month_translation_file = base_dir / "month_translation.json" # New Path object for this file

logger.info(f"Configured file paths:")
logger.info(f"  BUFR file: {bufr_file}")
logger.info(f"  24h Excel template: {xlsx_24h_template_path}")
logger.info(f"  Agri Excel template: {xlsx_agri_template_path}")
logger.info(f"  Output 24h Excel: {output_xlsx_24h_path}")
logger.info(f"  Output Agri Excel: {output_xlsx_agri_path}")
logger.info(f"  Word template: {word_template_path}")
logger.info(f"  Output Word: {output_word_path}")
logger.info(f"  Station mapping file: {mapping_file}")
logger.info(f"  Missing stations output: {missing_stations_path}")
logger.info(f"  Month translation file: {month_translation_file}")


# --- 4. Initial File Existence Checks (using pytaps.file_operations) ---
# Check if the synop BUFR file exists
if not check_file_exists_and_log(bufr_file, logger_instance=logger):
    logger.critical(f"SYNOP BUFR file not found at {bufr_file}. Exiting.")
    exit(1)

# --- 5. Load Station Mapping (with helper for JSON) ---
def _load_json_mapping(file_path: Path, logger_instance: logging.Logger) -> Dict[str, str]:
    """Helper to load JSON mapping with robust error handling."""
    if not check_file_exists_and_log(file_path, logger_instance=logger_instance):
        logger_instance.critical(f"Mapping file not found at {file_path}. Exiting.")
        exit(1)
    try:
        with open(file_path, 'r') as file:
            mapping = json.load(file)
        logger_instance.info(f"Mapping loaded successfully from {file_path}.")
        return mapping
    except json.JSONDecodeError:
        logger_instance.critical(f"Error decoding JSON from {file_path}. Check file format. Exiting.")
        exit(1)
    except Exception as e:
        logger_instance.critical(f"An unexpected error occurred while loading mapping from {file_path}: {e}. Exiting.", exc_info=True)
        exit(1)

station_mapping = _load_json_mapping(mapping_file, logger)

# Define today's date
today = datetime.now()
logger.info(f"Current date: {today.strftime('%Y-%m-%d %H:%M:%S')}")

# Check if today is the start of a new agricultural year (September 1)
reset_cumul = (today.month == 9 and today.day == 1)
if reset_cumul:
    logger.info("Today is September 1st. Cumulative precipitation will be reset.")
else:
    logger.info("Cumulative precipitation will not be reset today.")

# Define a red fill for cells where station data is missing
missing_station_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
logger.debug("Defined red fill for missing stations.")

# Missing stations list (collected for 24h precip)
missing_stations_24h = []

# --- 6. Read BUFR File (using pytaps.data_utils.read_bufr_to_dataframe) ---
df_06 = read_bufr_to_dataframe(
    bufr_file_path=bufr_file,
    columns=("stationOrSiteName", "totalPrecipitationOrTotalWaterEquivalent", "timePeriod"),
    logger_instance=logger
)
if df_06 is None: # read_bufr_to_dataframe returns None on error or empty
    logger.critical(f"Failed to read or process BUFR file {bufr_file}. Exiting.")
    exit(1)

# Rename columns for clarity
df_06.rename(columns={
    "stationOrSiteName": "SYNOP_STATION_NAME",
    "totalPrecipitationOrTotalWaterEquivalent": "Précip (mm)",
    "timePeriod": "Time"
}, inplace=True)
logger.info("Columns renamed in DataFrame.")
logger.debug(f"DataFrame columns after rename: {df_06.columns.tolist()}")

# Filter for 24-hour precipitation (Time == -24)
initial_rows = df_06.shape[0]
df_24h_precip = df_06[df_06["Time"] == -24].copy() # Use .copy() to avoid SettingWithCopyWarning
logger.info(f"Filtered DataFrame for 24-hour precipitation (Time == -24). Rows before filter: {initial_rows}, Rows after filter: {df_24h_precip.shape[0]}")

# Drop rows where stations are missing
initial_rows_dropna = df_24h_precip.shape[0]
df_24h_precip.dropna(subset=["SYNOP_STATION_NAME"], inplace=True)
logger.info(f"Dropped rows with missing 'SYNOP_STATION_NAME'. Rows before dropna: {initial_rows_dropna}, Rows after dropna: {df_24h_precip.shape[0]}")

df_24h_precip = df_24h_precip[["SYNOP_STATION_NAME", "Précip (mm)"]]
logger.debug(f"Selected columns 'SYNOP_STATION_NAME' and 'Précip (mm)'.")

# Sort DataFrame by station names
df_24h_precip.sort_values(["SYNOP_STATION_NAME"], inplace=True)
logger.info("DataFrame sorted by 'SYNOP_STATION_NAME'.")
logger.debug(f"Sample of 24h precipitation data:\n{df_24h_precip.head()}")

# --- Helper function for formatting precipitation values ---
def format_precip(value: Union[float, int, str]) -> str:
    """Formats precipitation value for display in Excel/Word."""
    if pd.isna(value) or (isinstance(value, (int, float)) and value <= 0):
        return '0'
    elif isinstance(value, (int, float)) and value < 0.1:
        return 'Tr'
    else:
        try:
            return f"{float(value):.1f}"
        except (ValueError, TypeError):
            logger.warning(f"Could not format value '{value}' to float. Returning as is.")
            return str(value)

logger.debug("Defined format_precip function.")

# --- 7. Update 24-Hour Precipitation Excel Template (using pytaps.data_utils.load_excel_workbook) ---
logger.info("Starting update of 24-hour precipitation Excel file.")
wb_24h, ws_24h = load_excel_workbook(xlsx_24h_template_path, logger_instance=logger)
# Error handling for load_excel_workbook is inside the function, it raises exceptions.

# --- Helper for updating Excel cells (specific to this script's structure) ---
def _update_excel_cells(
    worksheet: Any, # openpyxl.worksheet.worksheet.Worksheet
    data_df: pd.DataFrame,
    station_map: Dict[str, str],
    fill_style: PatternFill,
    logger_instance: logging.Logger,
    is_cumulative_update: bool = False,
    missing_stations_list: Optional[List[str]] = None # Only for 24h precip
) -> None:
    """
    Updates precipitation values in an Excel worksheet.
    Handles both 24h and cumulative updates.
    """
    for row_idx, row in enumerate(worksheet.iter_rows(min_row=2), start=2):
        station_template_names = [row[0].value, row[2].value, row[4].value] # Column A, C, E
        target_cells = [row[1], row[3], row[5]] # Column B, D, F

        for i, (template_name, target_cell) in enumerate(zip(station_template_names, target_cells)):
            if template_name is None:
                continue

            synop_name = station_map.get(template_name)
            if synop_name:
                station_data = data_df[data_df['SYNOP_STATION_NAME'] == synop_name]
                if not station_data.empty:
                    precip_value = station_data['Précip (mm)'].values[0]
                    if is_cumulative_update:
                        try:
                            current_cumul = float(target_cell.value) if target_cell.value is not None else 0.0
                        except (ValueError, TypeError):
                            logger_instance.warning(f"Row {row_idx}, Station '{template_name}': Invalid cumulative value '{target_cell.value}'. Treating as 0.0.")
                            current_cumul = 0.0
                        new_cumul = current_cumul + precip_value
                        target_cell.value = new_cumul
                        logger_instance.debug(f"Row {row_idx}, Station '{template_name}' ({synop_name}): Updated cumulative precip from {current_cumul:.1f} to {new_cumul:.1f} (added {precip_value:.1f})")
                    else:
                        target_cell.value = format_precip(precip_value)
                        logger_instance.debug(f"Row {row_idx}, Station '{template_name}' ({synop_name}): Updated 24h precip to {target_cell.value}")
                else:
                    # Data not found for the SYNOP station
                    if not is_cumulative_update:
                        target_cell.value = '/'
                        if missing_stations_list is not None:
                            missing_stations_list.append(template_name)
                        logger_instance.warning(f"Row {row_idx}, Station '{template_name}' ({synop_name}): Data not found in BUFR for 24h precip. Marked as missing.")
                    else:
                        # For cumulative, keep old value if no new data, but mark as missing
                        logger_instance.warning(f"Row {row_idx}, Station '{template_name}' ({synop_name}): Data not found in BUFR for cumulative update. Kept previous value.")
                    target_cell.fill = fill_style
            else:
                # No SYNOP mapping found for the template station name
                if not is_cumulative_update:
                    target_cell.value = '/'
                    if missing_stations_list is not None:
                        missing_stations_list.append(template_name)
                    logger_instance.warning(f"Row {row_idx}, Station '{template_name}': No SYNOP mapping found. Marked as missing for 24h precip.")
                else:
                    # For cumulative, keep old value if no mapping, but mark as missing
                    logger_instance.warning(f"Row {row_idx}, Station '{template_name}': No SYNOP mapping found for cumulative update. Kept previous value.")
                target_cell.fill = fill_style

_update_excel_cells(
    worksheet=ws_24h,
    data_df=df_24h_precip,
    station_map=station_mapping,
    fill_style=missing_station_fill,
    logger_instance=logger,
    is_cumulative_update=False,
    missing_stations_list=missing_stations_24h # Pass the list to collect missing stations
)
logger.info("Finished updating 24-hour precipitation Excel file.")

# Ensure parent directory exists before saving
ensure_parent_directory_exists(output_xlsx_24h_path, logger_instance=logger)
try:
    wb_24h.save(output_xlsx_24h_path)
    logger.info(f"Updated 24-hour precipitation Excel saved to {output_xlsx_24h_path}")
except Exception as e:
    logger.error(f"Error saving 24-hour precipitation Excel to {output_xlsx_24h_path}: {e}", exc_info=True)

# Save missing stations to a text file
if missing_stations_24h:
    try:
        unique_missing_stations = sorted(list(set(missing_stations_24h)))
        ensure_parent_directory_exists(missing_stations_path, logger_instance=logger)
        with open(missing_stations_path, 'w') as f:
            for station in unique_missing_stations:
                f.write(f"{station}\n")
        logger.warning(f"Found {len(unique_missing_stations)} unique missing stations. List saved in {missing_stations_path}")
    except Exception as e:
        logger.error(f"Error saving missing stations to {missing_stations_path}: {e}", exc_info=True)
else:
    logger.info("No missing stations found for 24-hour precipitation.")


# --- 8. Read and Update Cumulative Precipitation Template ---
logger.info("Starting update of cumulative precipitation Excel file.")
wb_agri, ws_agri = load_excel_workbook(xlsx_agri_template_path, logger_instance=logger)
# Error handling for load_excel_workbook is inside the function, it raises exceptions.

# Reset cumulative precipitation if today is September 1
if reset_cumul:
    logger.info("Resetting cumulative precipitation values in the template.")
    for row in ws_agri.iter_rows(min_row=2):
        for col_idx in [1, 3, 5]: # B, D, F
            if row[col_idx].value is not None:
                row[col_idx].value = 0
    logger.info("Cumulative precipitation values reset.")

# Add 24-hour precipitation to cumulative precipitation for each station
_update_excel_cells(
    worksheet=ws_agri,
    data_df=df_24h_precip, # Still using df_24h_precip for the added value
    station_map=station_mapping,
    fill_style=missing_station_fill,
    logger_instance=logger,
    is_cumulative_update=True,
    missing_stations_list=None # Not collecting missing stations here
)
logger.info("Finished updating cumulative precipitation Excel file.")

# Ensure parent directory exists before saving
ensure_parent_directory_exists(output_xlsx_agri_path, logger_instance=logger)
try:
    wb_agri.save(output_xlsx_agri_path)
    logger.info(f"Updated cumulative precipitation Excel saved to {output_xlsx_agri_path}")
except Exception as e:
    logger.error(f"Error saving cumulative precipitation Excel to {output_xlsx_agri_path}: {e}", exc_info=True)


# --- 9. Load Excel Data for Word doc (using pytaps.data_utils.load_excel_workbook) ---
logger.info(f"Loading output 24-hour precipitation Excel for Word document from {output_xlsx_24h_path}.")
wb_24h_word, ws_24h_word = load_excel_workbook(output_xlsx_24h_path, logger_instance=logger)

logger.info(f"Loading output cumulative precipitation Excel for Word document from {output_xlsx_agri_path}.")
wb_agri_word, ws_agri_word = load_excel_workbook(output_xlsx_agri_path, logger_instance=logger)

# --- 10. Update Word Template with 24-Hour and Cumulative Precipitation Tables ---
if not check_file_exists_and_log(word_template_path, logger_instance=logger):
    logger.critical(f"Word template not found at {word_template_path}. Exiting.")
    exit(1)
try:
    logger.info(f"Loading Word template from {word_template_path}.")
    document = Document(word_template_path)
    logger.info("Word template loaded successfully.")
except Exception as e:
    logger.critical(f"Error loading Word template: {e}. Exiting.", exc_info=True)
    exit(1)

# Load the month translation mapping (using helper)
month_mapping = _load_json_mapping(month_translation_file, logger)

# Generate the dynamic date lines
end_date_dt = datetime.now()
start_date_dt = datetime.now() - timedelta(days=1)

end_date_str = end_date_dt.strftime("%d %B %Y")
start_date_str = start_date_dt.strftime("%d %B %Y")

# Replace English month names with French equivalents
for english, french in month_mapping.items():
    end_date_str = end_date_str.replace(english, french)
    start_date_str = start_date_str.replace(english, french)

date_line_text = f"Quantités de pluie enregistrée du {start_date_str} à 06h au {end_date_str} à 06h."

# Determine the start year for the agricultural cumulative period
agri_start_year = end_date_dt.year
if end_date_dt.month < 9: # If current month is before September, the agri year started last year
    agri_start_year -= 1
agric_text = f"Cumuls de pluie enregistrés Validité : du 01 Septembre {agri_start_year} au {end_date_str}"
logger.info(f"Generated 24h date line: '{date_line_text}'")
logger.info(f"Generated cumulative date line: '{agric_text}'")

# Find the table in the Word document (assuming it's the first table)
try:
    table_24h_header = document.tables[0] # Header for 24h table
    table_24h_data = document.tables[1]   # Data table for 24h
    table_agri_header = document.tables[2] # Header for cumulative table
    table_agri_data = document.tables[3]   # Data table for cumulative
    logger.info("Accessed main tables (0, 1, 2, 3) in the Word document.")
except IndexError:
    logger.critical("Could not find expected tables in the Word document. Check template structure. Exiting.")
    exit(1)
except Exception as e:
    logger.critical(f"An unexpected error occurred while accessing Word tables: {e}. Exiting.", exc_info=True)
    exit(1)

# --- Helper for updating Word table headers ---
def _update_word_table_header(
    cell: Any, # docx.table._Cell
    text: str,
    font_size: Pt,
    font_color: RGBColor,
    bold: bool,
    logger_instance: logging.Logger
) -> None:
    """Updates a Word table cell with formatted text."""
    paragraph = cell.paragraphs[0]
    paragraph.clear() # Clear existing content
    paragraph.add_run(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = font_size
    run.font.color.rgb = font_color
    run.font.bold = bold
    logger_instance.debug(f"Updated Word table header with text: '{text}'")

_update_word_table_header(table_24h_header.cell(0, 0), date_line_text, Pt(14), RGBColor(0x00, 0x99, 0xDA), True, logger)
_update_word_table_header(table_agri_header.cell(0, 0), agric_text, Pt(14), RGBColor(0x00, 0x99, 0xDA), True, logger)


# --- Helper for populating Word data tables ---
def _populate_word_data_table(
    word_table: Any, # docx.table.Table
    excel_worksheet: Any, # openpyxl.worksheet.worksheet.Worksheet
    logger_instance: logging.Logger,
    font_size: Pt = Pt(9)
) -> None:
    """
    Populates a Word table from an Excel worksheet (columns B, D, F),
    and sets the font size for all cells in the table.
    """
    for row_index, excel_row in enumerate(excel_worksheet.iter_rows(min_row=2, values_only=True), start=1):
        for col_idx_word, excel_col_index in enumerate([1, 3, 5]): # Excel columns B, D, F
            value = format_precip(excel_row[excel_col_index])
            # Map Excel data column index (0, 1, 2 for B, D, F) to Word table column index (1, 3, 5)
            # Word table has 6 columns: (Station Name, Value, Station Name, Value, Station Name, Value)
            # So, Excel column 1 (B) goes to Word col 1. Excel col 3 (D) goes to Word col 3. Excel col 5 (F) goes to Word col 5.
            # This is `col_idx_word * 2 + 1` where `col_idx_word` is 0, 1, 2
            word_target_col = col_idx_word * 2 + 1

            if row_index < len(word_table.rows) and word_target_col < len(word_table.columns):
                cell = word_table.cell(row_index, word_target_col)
                # Clear existing content before adding new run
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(value)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = font_size # Apply font size directly to the run
                logger_instance.debug(f"Word Table: Row {row_index}, Col {word_target_col} updated with '{value}'.")
            else:
                logger_instance.warning(f"Word Table: Skipping row {row_index} or column {word_target_col} as it exceeds table dimensions.")

    # Ensure all cells in the table have the specified font size
    for row in word_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = font_size
    logger_instance.info(f"Finished populating Word table and set font size.")


logger.info("Populating 24-hour precipitation table in Word document.")
_populate_word_data_table(table_24h_data, ws_24h_word, logger)

logger.info("Populating cumulative precipitation table in Word document.")
_populate_word_data_table(table_agri_data, ws_agri_word, logger)


# Save the final Word document
ensure_parent_directory_exists(output_word_path, logger_instance=logger)
try:
    document.save(output_word_path)
    logger.info(f"Combined Word document saved successfully as {output_word_path}")
except Exception as e:
    logger.error(f"Error saving Word document to {output_word_path}: {e}", exc_info=True)

logger.info("Script execution finished.")
