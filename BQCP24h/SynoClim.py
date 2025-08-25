import os
import json
from datetime import datetime, timedelta
import pdbufr
import pandas as pd
from openpyxl import load_workbook # Still needed for the specific Excel template modification
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from pathlib import Path

# Import pytaps utilities
from pytaps.logging_utils import setup_logger
from pytaps.file_operations import check_file_exists_and_log
from pytaps.data_utils import load_excel_workbook # For loading the Excel template

import argparse
import logging

parser = argparse.ArgumentParser(description=f"{os.path.basename(__file__)} script.")

# Add positional arguments passed from Bulletin.py
parser.add_argument('year', type=str, help='Year (YYYY) for data processing.')
parser.add_argument('month', type=str, help='Month (MM) for data processing.')
parser.add_argument('day', type=str, help='Day (DD) for data processing.')
parser.add_argument('local_directory', type=str, help='Local base directory for file operations (equivalent to PWD).')


try:
    args = parser.parse_args()

except SystemExit as e:
    print(f"ERROR ({os.path.basename(__file__)}): Argument parsing failed. Exit code: {e.code}", file=sys.stderr)
    sys.exit(e.code)

shared_log_file_from_env = os.environ.get('PYTAPS_SHARED_LOG_FILE')

script_name_for_log = os.path.basename(__file__).replace(".py", "")
logger, actual_log_path_from_pytaps = setup_logger(
    script_name=script_name_for_log,
    log_directory_base=os.path.dirname(os.path.abspath(__file__)),
    log_level=logging.INFO,
    shared_log_file_path=shared_log_file_from_env
)

# Get values from command-line arguments (using args object)
AA = args.year
MM = args.month
DD = args.day
LOCAL_DIR = Path(args.local_directory) # Convert to Path object immediately

logger.info(f"Script {os.path.basename(__file__)} started with args: Year={AA}, Month={MM}, Day={DD}, LocalDir={LOCAL_DIR}")

# File paths - Convert to Path objects for better handling with pytaps functions
bufr_file = LOCAL_DIR / "Synop" / f"Synop_{AA}{MM}{DD}0600.bufr"
xlsx_template_path = LOCAL_DIR / "templates" / "template.xlsx"
output_xlsx_path = LOCAL_DIR / f"synop_precip_output{AA}{DD}0600.xlsx"
word_template_path = LOCAL_DIR / "templates" / "doc_template.docx"
output_word_path = LOCAL_DIR / f"SynopClim_precip_table{MM}{DD}0600.docx"
missing_stations_path = LOCAL_DIR / f"{AA}{MM}{DD}.txt"
json_mapping_path = LOCAL_DIR / "climaticstations.json"
excel_clim_data_path = LOCAL_DIR / "Climxlsx" / f"data_{AA}-{MM}-{DD}.xlsx" # Path for the additional Excel file
month_translation_path = LOCAL_DIR / "month_translation.json" # Assuming it's in LOCAL_DIR

logger.info(f"Configured file paths:")
logger.info(f"  BUFR file: {bufr_file}")
logger.info(f"  Excel template: {xlsx_template_path}")
logger.info(f"  Output Excel: {output_xlsx_path}")
logger.info(f"  Word template: {word_template_path}")
logger.info(f"  Output Word: {output_word_path}")
logger.info(f"  Missing stations file: {missing_stations_path}")
logger.info(f"  JSON mapping file: {json_mapping_path}")
logger.info(f"  Climatic Excel data: {excel_clim_data_path}")
logger.info(f"  Month translation JSON: {month_translation_path}")


# --- BUFR Reading and Processing ---
df_06 = pd.DataFrame() # Initialize to empty DataFrame
# Use check_file_exists_and_log before attempting to read
if not check_file_exists_and_log(bufr_file, logger_instance=logger):
    logger.critical(f"BUFR file not found at {bufr_file}. Exiting.")
    sys.exit(1)

try:
    logger.info(f"Attempting to read BUFR file: {bufr_file}")
    # Keep pdbufr.read_bufr as the pytaps function is specialized for temperature and doesn't fit this specific precipitation filtering logic.
    df_06 = pdbufr.read_bufr(bufr_file, columns=("stationOrSiteName",
                                                "totalPrecipitationOrTotalWaterEquivalent",
                                                "timePeriod"))
    logger.info(f"Successfully read BUFR file. Initial DataFrame shape: {df_06.shape}")
    logger.debug(f"First 5 rows of df_06:\n{df_06.head()}")
except Exception as e:
    logger.exception(f"An error occurred while reading the BUFR file: {e}. Exiting.")
    sys.exit(1)

# Rename columns for clarity
logger.info("Renaming DataFrame columns.")
df_06.rename(columns={"stationOrSiteName": "STATIONS",
                        "totalPrecipitationOrTotalWaterEquivalent": "Précip (mm)",
                        "timePeriod": "Time"}, inplace=True)
logger.debug(f"DataFrame columns after rename: {df_06.columns.tolist()}")

# Filter for 24-hour precipitation (Time == -24)
logger.info("Filtering for 24-hour precipitation (Time == -24).")
df_24h_precip = df_06[df_06["Time"] == -24]
logger.info(f"Filtered 24-hour precipitation DataFrame shape: {df_24h_precip.shape}")
logger.debug(f"First 5 rows of df_24h_precip:\n{df_24h_precip.head()}")

# Drop rows where station names are missing
logger.info("Dropping rows with missing station names.")
initial_shape = df_24h_precip.shape[0]
df_24h_precip = df_24h_precip.dropna(subset=["STATIONS"])
logger.info(f"Dropped {initial_shape - df_24h_precip.shape[0]} rows with missing station names. New shape: {df_24h_precip.shape}")
df_24h_precip = df_24h_precip[["STATIONS", "Précip (mm)"]]
logger.debug(f"Final df_24h_precip columns: {df_24h_precip.columns.tolist()}")

# --- Load the Excel template using PyTAP ---
wb = None
try:
    logger.info(f"Attempting to load Excel template: {xlsx_template_path}")
    wb, ws = load_excel_workbook(xlsx_template_path, logger_instance=logger)
    logger.info("Successfully loaded Excel template.")
except FileNotFoundError: # load_excel_workbook already logs, but we need to exit here
    logger.critical(f"Excel template file not found: {xlsx_template_path}. Exiting.")
    sys.exit(1)
except Exception as e: # load_excel_workbook already logs, but we need to exit here
    logger.critical(f"An error occurred while loading the Excel template: {e}. Exiting.", exc_info=True)
    sys.exit(1)

# --- Load the JSON mapping for station names ---
station_mapping = {}
if not check_file_exists_and_log(json_mapping_path, logger_instance=logger):
    logger.critical(f"JSON mapping file not found: {json_mapping_path}. This file is required. Exiting.")
    sys.exit(1)

try:
    logger.info(f"Attempting to load JSON mapping file: {json_mapping_path}")
    with open(json_mapping_path, 'r') as f:
        station_mapping = json.load(f)
    logger.info(f"Successfully loaded JSON mapping. {len(station_mapping)} entries found.")
    logger.debug(f"Sample of station_mapping: {list(station_mapping.items())[:5]}")
except json.JSONDecodeError:
    logger.critical(f"Error decoding JSON from {json_mapping_path}. Please check file format. Exiting.")
    sys.exit(1)
except Exception as e:
    logger.exception(f"An error occurred while loading JSON mapping: {e}. Exiting.")
    sys.exit(1)

# Missing stations list
missing_stations = []

# Function to format precipitation values (Application-specific, keep as is)
def format_precip(value):
    if pd.isna(value) or value <= 0:
        return '0'
    elif value < 0.1:
        return 'Tr'
    else:
        return f"{value:.1f}"  # Format to one decimal place

# Check if the additional Excel file exists and load it
excel_precip_data = None
if check_file_exists_and_log(excel_clim_data_path, logger_instance=logger):
    try:
        logger.info(f"Additional Excel file found: {excel_clim_data_path}. Attempting to load.")
        # Keeping pd.read_excel as pytaps.data_utils.load_excel_workbook returns openpyxl objects, not a DataFrame.
        excel_precip_data = pd.read_excel(excel_clim_data_path)
        # Ensure that the column names match your file structure
        excel_precip_data.columns = [col.strip() for col in excel_precip_data.columns]  # Strip any whitespace
        logger.info(f"Successfully loaded additional Excel file. Columns: {excel_precip_data.columns.tolist()}")
        if 'Station' not in excel_precip_data.columns or 'Precipitation' not in excel_precip_data.columns:
            logger.warning(f"Expected columns 'Station' and 'Precipitation' not found in {excel_clim_data_path}. Data might not be correctly processed from this file.")
        logger.debug(f"First 5 rows of additional Excel data:\n{excel_precip_data.head()}")
    except Exception as e:
        logger.exception(f"An error occurred while reading the additional Excel file {excel_clim_data_path}: {e}. This file will not be used.")
        excel_precip_data = None # Ensure it's None if loading fails
else:
    logger.info(f"Additional Excel file not found: {excel_clim_data_path}. Skipping climatic station data from this source.")

# Function to get precipitation data from the Excel file for climatic stations (Application-specific, keep as is)
def get_precip_from_excel(station_name):
    if excel_precip_data is not None:
        mapped_name = station_mapping.get(station_name, station_name)  # Get mapped name or use original if not found
        logger.debug(f"Searching for station '{station_name}' (mapped to '{mapped_name}') in additional Excel data.")
        # Check if 'Station' and 'Precipitation' columns exist before trying to access them
        if 'Station' in excel_precip_data.columns and 'Precipitation' in excel_precip_data.columns:
            precip_row = excel_precip_data[excel_precip_data['Station'] == mapped_name]
            if not precip_row.empty:
                precip_value = precip_row['Precipitation'].values[0]
                logger.debug(f"Found precipitation {precip_value} for '{station_name}' in additional Excel.")
                return precip_value
            else:
                logger.debug(f"Station '{station_name}' (mapped to '{mapped_name}') not found in additional Excel data.")
        else:
            logger.warning(f"Cannot get precipitation from additional Excel: 'Station' or 'Precipitation' column missing.")
    return None

logger.info("Starting to loop through Excel template rows to populate precipitation values.")
# Loop through rows in the Excel template to populate values (Application-specific, keep as is)
for row_idx, row in enumerate(ws.iter_rows(min_row=3), start=3):  # Start from the first data row after headers
    station_1_cell = row[1]
    station_2_cell = row[4]

    station_1 = station_1_cell.value  # Station name in the 2nd column (index 1)
    station_2 = station_2_cell.value  # Station name in the 5th column (index 4)

    # Process station_1
    if station_1:
        logger.debug(f"Processing row {row_idx}, Station 1: '{station_1}'")
        precip_value_1 = None
        # Check if station is in the BUFR data
        station_data_1 = df_24h_precip[df_24h_precip['STATIONS'] == station_1]
        if not station_data_1.empty:
            precip_value_1 = station_data_1['Précip (mm)'].values[0]
            row[2].value = format_precip(precip_value_1)  # Place the precipitation value in the 3rd column (index 2)
            logger.debug(f"  Found BUFR precip for '{station_1}': {precip_value_1}. Formatted: {row[2].value}")
        else:
            logger.debug(f"  Station '{station_1}' not found in BUFR data. Checking additional Excel.")
            if excel_precip_data is not None:
                # Check if station is in the Excel data
                precip_value_1 = get_precip_from_excel(station_1)
                if precip_value_1 is not None:
                    row[2].value = format_precip(precip_value_1)
                    logger.debug(f"  Found Excel precip for '{station_1}': {precip_value_1}. Formatted: {row[2].value}")
                else:
                    row[2].value = '/'  # Mark as missing
                    missing_stations.append(station_1)  # Add to missing list
                    logger.info(f"  Station '{station_1}' not found in any data source. Marked as missing ('/').")
            else:
                row[2].value = '/'  # Mark as missing
                missing_stations.append(station_1)  # Add to missing list
                logger.info(f"  Station '{station_1}' not found in BUFR data and no additional Excel data available. Marked as missing ('/').")
    else:
        logger.debug(f"Row {row_idx}, Station 1 cell is empty. Skipping.")

    # Process station_2
    if station_2:
        logger.debug(f"Processing row {row_idx}, Station 2: '{station_2}'")
        precip_value_2 = None
        # Check if station is in the BUFR data
        station_data_2 = df_24h_precip[df_24h_precip['STATIONS'] == station_2]
        if not station_data_2.empty:
            precip_value_2 = station_data_2['Précip (mm)'].values[0]
            row[5].value = format_precip(precip_value_2)  # Place the precipitation value in the 6th column (index 5)
            logger.debug(f"  Found BUFR precip for '{station_2}': {precip_value_2}. Formatted: {row[5].value}")
        else:
            logger.debug(f"  Station '{station_2}' not found in BUFR data. Checking additional Excel.")
            if excel_precip_data is not None:
                # Check if station is in the Excel data
                precip_value_2 = get_precip_from_excel(station_2)
                if precip_value_2 is not None:
                    row[5].value = format_precip(precip_value_2)
                    logger.debug(f"  Found Excel precip for '{station_2}': {precip_value_2}. Formatted: {row[5].value}")
                else:
                    row[5].value = '/'  # Mark as missing
                    missing_stations.append(station_2)  # Add to missing list
                    logger.info(f"  Station '{station_2}' not found in any data source. Marked as missing ('/').")
            else:
                row[5].value = '/'  # Mark as missing
                missing_stations.append(station_2)  # Add to missing list
                logger.info(f"  Station '{station_2}' not found in BUFR data and no additional Excel data available. Marked as missing ('/').")
    else:
        logger.debug(f"Row {row_idx}, Station 2 cell is empty. Skipping.")

logger.info(f"Finished populating Excel template. Total missing stations identified: {len(missing_stations)}")
logger.debug(f"Missing stations list: {missing_stations}")

# Save the updated Excel file (Keeping original wb.save as pytaps save_dataframe_to_excel creates a new file from DataFrame)
try:
    logger.info(f"Attempting to save updated Excel file to: {output_xlsx_path}")
    wb.save(output_xlsx_path)
    logger.info(f"Precipitation values successfully inserted into {output_xlsx_path}")
except Exception as e:
    logger.exception(f"An error occurred while saving the output Excel file: {e}. Please check file permissions or if the file is open.")

# Save missing stations to a text file (Application-specific, keep as is)
try:
    logger.info(f"Attempting to save missing stations list to: {missing_stations_path}")
    # Ensure parent directory exists for the missing stations file
    missing_stations_path.parent.mkdir(parents=True, exist_ok=True)
    with open(missing_stations_path, 'w') as f:
        for station in missing_stations:
            f.write(f"{station}\n")
    logger.info(f"Missing stations saved in {missing_stations_path}")
except Exception as e:
    logger.exception(f"An error occurred while saving missing stations to file: {e}.")

# --- Copying the output into a Word file ---
doc = None
if not check_file_exists_and_log(word_template_path, logger_instance=logger):
    logger.critical(f"Word template file not found: {word_template_path}. Exiting.")
    sys.exit(1)
try:
    logger.info(f"Attempting to load Word template: {word_template_path}")
    doc = Document(word_template_path)
    logger.info("Successfully loaded Word template.")
except Exception as e:
    logger.exception(f"An error occurred while loading the Word template: {e}. Exiting.")
    sys.exit(1)

# Load month translation JSON
month_translation = {}
if not check_file_exists_and_log(month_translation_path, logger_instance=logger):
    logger.error("month_translation.json not found. Date line will not be translated. Please ensure this file exists.")
else:
    try:
        logger.info(f"Attempting to load month_translation.json from {month_translation_path}.")
        with open(month_translation_path, 'r') as f:
            month_translation = json.load(f)["months"]
        logger.info("Successfully loaded month translation JSON.")
    except json.JSONDecodeError:
        logger.error(f"Error decoding JSON from {month_translation_path}. Date line will not be translated.")
    except KeyError:
        logger.error(f"Key 'months' not found in {month_translation_path}. Date line will not be translated.")
    except Exception as e:
        logger.exception(f"An unexpected error occurred while loading {month_translation_path}: {e}.")


# Function to translate months (Application-specific, keep as is)
def translate_month(date_str):
    if not month_translation: # Skip if translation failed to load or is empty
        return date_str
    for eng_month, fr_month in month_translation.items():
        if eng_month in date_str:
            return date_str.replace(eng_month, fr_month)
    return date_str

# Generate the dynamic date line
end_date = translate_month(datetime.now().strftime("%d %B %Y"))
start_date = translate_month((datetime.now() - timedelta(days=1)).strftime("%d %B %Y"))
date_line_text = f"Quantités de pluie enregistrée du {start_date} à 06h au {end_date} à 06h."
logger.info(f"Generated date line for Word document: '{date_line_text}'")

try:
    # Find the table in the Word document (assuming it's the first table)
    logger.info("Attempting to update the first table (date line) in the Word document.")
    table_date_line = doc.tables[0]

    # Insert the dynamic date line into the first cell of the table
    cell = table_date_line.cell(0, 0)  # Access the first row and first column
    cell.text = date_line_text

    # Format the text inside the table cell
    paragraph = cell.paragraphs[0]  # Get the first paragraph of the cell
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.runs[0]  # Access the first run of the paragraph
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x99, 0xDA)  # Set text color to #0099DA
    run.font.bold = True  # Make the text bold
    logger.debug("Date line inserted and formatted in Word document.")

    # Add spacing after the date line (optional)
    doc.add_paragraph()
    logger.debug("Added an empty paragraph after the date line.")

    # Find the table in the Word document (assuming it's the second table for data)
    logger.info("Attempting to populate the second table (precipitation data) in the Word document.")
    table_data = doc.tables[1]

    # Copy data from the Excel file into the Word table
    row_index_word = 1 # Start from the second row in Word table (index 1) as row 0 is header
    for row_idx_excel in range(3, ws.max_row + 1): # Iterate through data rows in Excel
        precip_1 = ws.cell(row=row_idx_excel, column=3).value
        precip_2 = ws.cell(row=row_idx_excel, column=6).value

        if row_index_word < len(table_data.rows):
            logger.debug(f"Populating Word table row {row_index_word} with Excel data from row {row_idx_excel}: Precip1='{precip_1}', Precip2='{precip_2}'")
            # Set the value for the first station (column 2 in Word table)
            cell_1 = table_data.cell(row_index_word, 2)
            cell_1.text = str(precip_1) if precip_1 is not None else '/'
            # Set font size to 9 for the first station
            for paragraph in cell_1.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

            # Set the value for the second station (column 5 in Word table)
            cell_2 = table_data.cell(row_index_word, 5)
            cell_2.text = str(precip_2) if precip_2 is not None else '/'
            # Set font size to 9 for the second station
            for paragraph in cell_2.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            row_index_word += 1
        else:
            logger.warning(f"Word table has fewer rows than Excel data. Skipping Excel row {row_idx_excel} (Precip1: {precip_1}, Precip2: {precip_2}) as no corresponding Word table row {row_index_word} exists.")

    logger.info("Finished populating Word document table.")

    # Save the filled Word document
    logger.info(f"Attempting to save filled Word document to: {output_word_path}")
    doc.save(output_word_path)
    logger.info(f"Precipitation values successfully inserted and formatted in {output_word_path}")

except IndexError as e:
    logger.error(f"Error accessing table or cell in Word document. Ensure the template has at least two tables and sufficient rows/columns: {e}", exc_info=True)
except Exception as e:
    logger.exception(f"An unexpected error occurred while processing or saving the Word document: {e}.")

logger.info(f"Script {os.path.basename(__file__)} finished.")
