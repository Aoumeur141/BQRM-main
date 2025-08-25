import json
import os # Import os for environment variables
import pandas as pd
from datetime import datetime, timedelta
from openpyxl.styles import PatternFill
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from pathlib import Path

# Import functions from pytaps
from pytaps.logging_utils import setup_logger
from pytaps.file_operations import check_file_exists_and_log
from pytaps.data_utils import load_excel_workbook, read_bufr_to_dataframe

import argparse
import sys

parser = argparse.ArgumentParser(description=f"{os.path.basename(__file__)} script.")


# Add positional arguments specific to this script (still needed)
parser.add_argument('year', type=str, help='Year (YYYY) for data processing.')
parser.add_argument('month', type=str, help='Month (MM) for data processing.')
parser.add_argument('day', type=str, help='Day (DD) for data processing.')
parser.add_argument('local_directory', type=str, help='Local base directory for file operations.')

# --- Debugging Prints for Argument Parsing ---
try:
    # Now parse_args() should work directly as --shared-log-file is no longer expected
    args = parser.parse_args() 
except SystemExit as e:
    print(f"ERROR ({os.path.basename(__file__)}): Argument parsing failed. Exit code: {e.code}", file=sys.stderr)
    sys.exit(e.code)

# --- NEW BLOCK: Get shared log file path from environment variable ---
shared_log_file_from_env = os.environ.get('PYTAPS_SHARED_LOG_FILE')
if shared_log_file_from_env:
    logger_setup_message = f"Shared log file path from environment: {shared_log_file_from_env}"
    # print(f"DEBUG ({os.path.basename(__file__)}): {logger_setup_message}", file=sys.stderr) # Keep for initial debugging, remove later
else:
    logger_setup_message = "PYTAPS_SHARED_LOG_FILE environment variable not set. This script will create its own log file."
    # print(f"WARNING ({os.path.basename(__file__)}): {logger_setup_message}", file=sys.stderr) # Keep for initial debugging, remove later
# --- END NEW BLOCK ---

script_name_for_log = os.path.basename(__file__).replace(".py", "")
logger, actual_log_path_from_pytaps = setup_logger(
    script_name=script_name_for_log,
    log_directory_base=os.path.dirname(os.path.abspath(__file__)),
    log_level=logging.INFO,
    shared_log_file_path=shared_log_file_from_env # Use the path from the environment variable
)
logger.info(logger_setup_message) # Log the message about source of log path


# Get values from command-line arguments (using args object)
AA = args.year
MM = args.month
DD = args.day
LOCAL_DIR = Path(args.local_directory) # Convert to Path object immediately

logger.info(f"Script {os.path.basename(__file__)} started with args: Year={AA}, Month={MM}, Day={DD}, LocalDir={LOCAL_DIR}")

# File paths with dynamic synop filename
bufr_file = LOCAL_DIR / "Synop" / f"Synop_{AA}{MM}{DD}0600.bufr"
xlsx_24h_template_path = LOCAL_DIR / "templates" / "cumul24.xlsx"
xlsx_agri_template_path = LOCAL_DIR / "templates" / "agricole.xlsx"
output_xlsx_24h_path = LOCAL_DIR / f"cumul24_{MM}{DD}0600.xlsx"
output_xlsx_agri_path = LOCAL_DIR / f"agricole_{MM}{DD}0600.xlsx"
word_template_path = LOCAL_DIR / "templates" / "cumul.docx"
output_word_path = LOCAL_DIR / f"Cumul_table{MM}{DD}0600.docx"
mapping_file = LOCAL_DIR / "ListStation.json"
missing_stations_path = LOCAL_DIR / f"Synop{AA}{MM}{DD}-06.txt"
month_translation_file = LOCAL_DIR / "month_translation.json" # Assuming it's in LOCAL_DIR

logger.info(f"Constructed file paths:")
logger.info(f"  BUFR file: {bufr_file}")
logger.info(f"  24h Excel template: {xlsx_24h_template_path}")
logger.info(f"  Agri Excel template: {xlsx_agri_template_path}")
logger.info(f"  Output 24h Excel: {output_xlsx_24h_path}")
logger.info(f"  Output Agri Excel: {output_xlsx_agri_path}")
logger.info(f"  Word template: {word_template_path}")
logger.info(f"  Output Word: {output_word_path}")
logger.info(f"  Station mapping file: {mapping_file}")
logger.info(f"  Missing stations log: {missing_stations_path}")
logger.info(f"  Month translation file: {month_translation_file}")


# Check if the synop BUFR file exists using pytaps utility
if not check_file_exists_and_log(bufr_file, logger_instance=logger):
    logger.critical(f"SYNOP BUFR file not found at {bufr_file}. Exiting script.")
    sys.exit(1)

# Load station name mapping from JSON
station_mapping = {}
if not check_file_exists_and_log(mapping_file, logger_instance=logger):
    logger.critical(f"Station mapping file not found at {mapping_file}. Exiting script.")
    sys.exit(1)

try:
    logger.info(f"Attempting to load station mapping from {mapping_file}...")
    with open(mapping_file, 'r') as file:
        station_mapping = json.load(file)
    logger.info(f"Successfully loaded station mapping from {mapping_file}. Found {len(station_mapping)} stations.")
except json.JSONDecodeError as e:
    logger.critical(f"Error decoding JSON from {mapping_file}: {e}. Exiting script.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"An unexpected error occurred while loading station mapping: {e}. Exiting script.", exc_info=True)
    sys.exit(1)

# Define today's date
today = datetime.now()
logger.info(f"Today's date: {today.strftime('%Y-%m-%d')}")

# Check if today is the start of a new agricultural year (September 1)
reset_cumul = (today.month == 9 and today.day == 1)
if reset_cumul:
    logger.info("Today is September 1st. Cumulative precipitation will be reset.")
else:
    logger.info("Cumulative precipitation will not be reset.")

# Define a red fill for cells where station data is missing
missing_station_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

# Missing stations list
missing_stations = []

# Read the BUFR file and extract relevant columns using pytaps utility
df_06 = pd.DataFrame() # Initialize as empty DataFrame
bufr_columns_to_read = ["stationOrSiteName", "totalPrecipitationOrTotalWaterEquivalent", "timePeriod"]
df_06 = read_bufr_to_dataframe(bufr_file, columns=bufr_columns_to_read, logger_instance=logger)

if df_06 is None or df_06.empty:
    logger.critical("Failed to load BUFR data or DataFrame is empty. Exiting script.")
    sys.exit(1)

# Rename columns for clarity
df_06.rename(columns={"stationOrSiteName": "SYNOP_STATION_NAME",
                        "totalPrecipitationOrTotalWaterEquivalent": "Précip (mm)",
                        "timePeriod": "Time"}, inplace=True)
logger.info("Columns renamed in DataFrame.")

# Filter for 24-hour precipitation (Time == -24)
df_24h_precip = df_06[df_06["Time"] == -24]
logger.info(f"Filtered for 24-hour precipitation (Time == -24). Records found: {df_24h_precip.shape[0]}")

# Drop rows where stations are missing
initial_rows = df_24h_precip.shape[0]
df_24h_precip = df_24h_precip.dropna(subset=["SYNOP_STATION_NAME"])
if df_24h_precip.shape[0] < initial_rows:
    logger.warning(f"Dropped {initial_rows - df_24h_precip.shape[0]} rows due to missing 'SYNOP_STATION_NAME' after filtering.")
df_24h_precip = df_24h_precip[["SYNOP_STATION_NAME", "Précip (mm)"]]
logger.info(f"Selected 'SYNOP_STATION_NAME' and 'Précip (mm)' columns. Current DataFrame shape: {df_24h_precip.shape}")

# Sort DataFrame by station names
df_24h_precip = df_24h_precip.sort_values(["SYNOP_STATION_NAME"])
logger.info("DataFrame sorted by 'SYNOP_STATION_NAME'.")

# --- Read 24-Hour Precipitation Template ---
# Use pytaps.data_utils.load_excel_workbook
wb_24h, ws_24h = None, None # Initialize to None
try:
    wb_24h, ws_24h = load_excel_workbook(xlsx_24h_template_path, logger_instance=logger)
except FileNotFoundError: # load_excel_workbook already logs, just need to exit
    logger.critical(f"24-hour precipitation Excel template not found at {xlsx_24h_template_path}. Exiting script.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading 24-hour precipitation Excel template from {xlsx_24h_template_path}: {e}. Exiting script.", exc_info=True)
    sys.exit(1)

def format_precip(value):
    """Formats precipitation value: '0' for <=0, 'Tr' for <0.1, else one decimal."""
    try:
        if pd.isna(value) or value <= 0:
            return '0'
        elif value < 0.1:
            return 'Tr'
        else:
            return f"{value:.1f}"
    except (TypeError, ValueError):
        logger.warning(f"Could not format precipitation value: '{value}'. Returning '/'.")
        return '/'

# --- Update 24-Hour Precipitation Template ---
logger.info("Updating 24-hour precipitation Excel file with data.")
for row_idx, row in enumerate(ws_24h.iter_rows(min_row=2), start=2): # Start from row 2
    # Get station names from the Excel template columns
    station_template_1 = row[0].value  # Column A
    station_template_2 = row[2].value  # Column C
    station_template_3 = row[4].value  # Column E

    # Helper function to process each station column
    def process_station_column(template_name, excel_value_col_idx):
        if template_name:
            synop_station_name = station_mapping.get(template_name)
            if synop_station_name:
                station_data = df_24h_precip[df_24h_precip['SYNOP_STATION_NAME'] == synop_station_name]
                if not station_data.empty:
                    precip_value = station_data['Précip (mm)'].values[0]
                    row[excel_value_col_idx].value = format_precip(precip_value)
                    logger.debug(f"Row {row_idx}: Station '{template_name}' ({synop_station_name}) found, precip: {precip_value}")
                else:
                    row[excel_value_col_idx].value = '/'
                    row[excel_value_col_idx].fill = missing_station_fill
                    missing_stations.append(template_name)
                    logger.warning(f"Row {row_idx}: Station '{template_name}' ({synop_station_name}) not found in BUFR data. Marked as missing.")
            else:
                row[excel_value_col_idx].value = '/'
                row[excel_value_col_idx].fill = missing_station_fill
                missing_stations.append(template_name)
                logger.warning(f"Row {row_idx}: Station '{template_name}' not found in mapping. Marked as missing.")
        else:
            logger.debug(f"Row {row_idx}: Template column {chr(65 + excel_value_col_idx-1)} is empty, skipping.")

    process_station_column(station_template_1, 1) # Value in B (index 1)
    process_station_column(station_template_2, 3) # Value in D (index 3)
    process_station_column(station_template_3, 5) # Value in F (index 5)

# Save the updated 24-hour precipitation Excel file
try:
    wb_24h.save(output_xlsx_24h_path)
    logger.info(f"Updated 24-hour precipitation Excel file saved to {output_xlsx_24h_path}")
except Exception as e:
    logger.error(f"Error saving 24-hour precipitation Excel file to {output_xlsx_24h_path}: {e}", exc_info=True)

# Save missing stations to a text file
if missing_stations:
    unique_missing_stations = sorted(list(set(missing_stations))) # Use set to avoid duplicates, then sort
    try:
        # Ensure parent directory exists for the missing stations file
        missing_stations_path.parent.mkdir(parents=True, exist_ok=True)
        with open(missing_stations_path, 'w') as f:
            for station in unique_missing_stations:
                f.write(f"{station}\n")
        logger.info(f"Missing stations list saved to {missing_stations_path}. Total unique missing stations: {len(unique_missing_stations)}")
    except Exception as e:
        logger.error(f"Error saving missing stations to {missing_stations_path}: {e}", exc_info=True)
else:
    logger.info("No missing stations found for 24-hour precipitation data.")


# --- Read and Update Cumulative Precipitation Template ---
# Use pytaps.data_utils.load_excel_workbook
wb_agri, ws_agri = None, None # Initialize to None
try:
    wb_agri, ws_agri = load_excel_workbook(xlsx_agri_template_path, logger_instance=logger)
except FileNotFoundError:
    logger.critical(f"Cumulative precipitation Excel template not found at {xlsx_agri_template_path}. Exiting script.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading cumulative precipitation Excel template from {xlsx_agri_template_path}: {e}. Exiting script.", exc_info=True)
    sys.exit(1)

# Reset cumulative precipitation if today is September 1
if reset_cumul:
    logger.info("Resetting cumulative precipitation values in the template.")
    for row in ws_agri.iter_rows(min_row=2): # Iterate over all cells that might contain cumulative values
        # Assuming cumulative values are in columns B, D, F (indices 1, 3, 5)
        # Ensure values are treated as numbers before resetting
        for col_idx in [1, 3, 5]:
            cell_value = row[col_idx].value
            if cell_value is not None and isinstance(cell_value, (int, float)):
                row[col_idx].value = 0
            elif cell_value is not None:
                logger.debug(f"Skipping reset for non-numeric cell at row {row.row}, col {col_idx+1}: '{cell_value}'")
    logger.info("Cumulative precipitation values reset to 0.")

# Add 24-hour precipitation to cumulative precipitation for each station
logger.info("Updating cumulative precipitation Excel file with daily data.")
for row_idx, row in enumerate(ws_agri.iter_rows(min_row=2), start=2):
    station_template_1 = row[0].value
    station_template_2 = row[2].value
    station_template_3 = row[4].value

    # Helper function to process each station column for cumulative update
    def process_cumulative_column(template_name, current_cumul_value_cell, excel_value_col_idx):
        if template_name:
            synop_station_name = station_mapping.get(template_name)
            if synop_station_name:
                station_data = df_24h_precip[df_24h_precip['SYNOP_STATION_NAME'] == synop_station_name]
                if not station_data.empty:
                    precip_value = station_data['Précip (mm)'].values[0]
                    # Ensure current_cumul_value is a number for arithmetic
                    try:
                        current_cumul_numeric = float(current_cumul_value_cell.value) if current_cumul_value_cell.value is not None else 0.0
                    except (ValueError, TypeError):
                        current_cumul_numeric = 0.0
                        logger.warning(f"Non-numeric cumulative value found for station '{template_name}' at row {row_idx}, col {excel_value_col_idx+1}. Treating as 0 for addition.")

                    new_cumul = current_cumul_numeric + precip_value
                    current_cumul_value_cell.value = new_cumul
                    logger.debug(f"Row {row_idx}: Station '{template_name}' ({synop_station_name}) cumulative updated from {current_cumul_numeric:.1f} to {new_cumul:.1f} (added {precip_value:.1f})")
                else:
                    current_cumul_value_cell.fill = missing_station_fill
                    logger.warning(f"Row {row_idx}: Station '{template_name}' ({synop_station_name}) not found in BUFR data for cumulative update. Cumulative unchanged.")
            else:
                current_cumul_value_cell.fill = missing_station_fill
                logger.warning(f"Row {row_idx}: Station '{template_name}' not found in mapping for cumulative update. Cumulative unchanged.")
        else:
            logger.debug(f"Row {row_idx}: Template column {chr(65 + excel_value_col_idx-1)} is empty, skipping cumulative update.")

    # Pass the cell objects directly to the helper function
    process_cumulative_column(station_template_1, row[1], 1)
    process_cumulative_column(station_template_2, row[3], 3)
    process_cumulative_column(station_template_3, row[5], 5)

# Save the updated cumulative precipitation Excel file
try:
    wb_agri.save(output_xlsx_agri_path)
    logger.info(f"Updated cumulative precipitation Excel file saved to {output_xlsx_agri_path}")
except Exception as e:
    logger.error(f"Error saving cumulative precipitation Excel file to {output_xlsx_agri_path}: {e}", exc_info=True)

# --- Load 24-Hour Precipitation Excel Data (for Word) ---
# Use pytaps.data_utils.load_excel_workbook
wb_24h_word, ws_24h_word = None, None # Initialize to None
try:
    wb_24h_word, ws_24h_word = load_excel_workbook(output_xlsx_24h_path, logger_instance=logger)
except FileNotFoundError:
    logger.critical(f"Output 24-hour precipitation Excel file not found at {output_xlsx_24h_path}. Cannot generate Word document. Exiting script.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading output 24-hour precipitation Excel file for Word from {output_xlsx_24h_path}: {e}. Cannot generate Word document. Exiting script.", exc_info=True)
    sys.exit(1)

# --- Load Cumulative Precipitation Excel Data (for Word) ---
# Use pytaps.data_utils.load_excel_workbook
wb_agri_word, ws_agri_word = None, None # Initialize to None
try:
    wb_agri_word, ws_agri_word = load_excel_workbook(output_xlsx_agri_path, logger_instance=logger)
except FileNotFoundError:
    logger.critical(f"Output cumulative precipitation Excel file not found at {output_xlsx_agri_path}. Cannot generate Word document. Exiting script.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading output cumulative precipitation Excel file for Word from {output_xlsx_agri_path}: {e}. Cannot generate Word document. Exiting script.", exc_info=True)
    sys.exit(1)

# --- Update Word Template with 24-Hour and Cumulative Precipitation Tables ---
document = None
if not check_file_exists_and_log(word_template_path, logger_instance=logger):
    logger.critical(f"Word template not found at {word_template_path}. Exiting script.")
    sys.exit(1)

try:
    logger.info(f"Loading Word template from {word_template_path}...")
    document = Document(word_template_path)
    logger.info("Word template loaded successfully.")
except Exception as e:
    logger.critical(f"Error loading Word template from {word_template_path}: {e}. Exiting script.", exc_info=True)
    sys.exit(1)

# Formatting function to ensure values have one decimal
def format_value(value):
    """Formats values for Word document: 'Tr', '0', '/' pass through, numbers to one decimal."""
    try:
        # Check if the value is already a string like 'Tr', '0', or '/' (from previous Excel steps)
        if isinstance(value, str) and value.strip().lower() in ['tr', '0', '/']:
            return value.strip()
        # Otherwise, try to convert to float and format
        return f"{float(value):.1f}"
    except (ValueError, TypeError):
        logger.warning(f"Could not format value for Word: '{value}'. Returning '/'.")
        return "/"

# Load the month translation mapping from JSON
month_mapping = {}
if not check_file_exists_and_log(month_translation_file, logger_instance=logger):
    logger.warning(f"Month translation file not found at {month_translation_file}. Using default English month names.")
else:
    try:
        logger.info(f"Attempting to load month translation from {month_translation_file}...")
        with open(month_translation_file, "r") as month_file:
            month_mapping = json.load(month_file)["months"]
        logger.info(f"Successfully loaded month translation from {month_translation_file}.")
    except json.JSONDecodeError as e:
        logger.error(f"Error decoding JSON from {month_translation_file}: {e}. Using default English month names.")
    except KeyError:
        logger.error(f"Key 'months' not found in {month_translation_file}. Using default English month names.")
    except Exception as e:
        logger.error(f"An unexpected error occurred while loading month translation: {e}. Using default English month names.", exc_info=True)

# Generate the dynamic date lines
end_date = datetime.now().strftime("%d %B %Y")
start_date = (datetime.now() - timedelta(days=1)).strftime("%d %B %Y")

# Replace English month names with French equivalents
for english, french in month_mapping.items():
    end_date = end_date.replace(english, french)
    start_date = start_date.replace(english, french)

date_line_text = f"Quantités de pluie enregistrée du {start_date} à 06h au {end_date} à 06h."
# Assuming AA is the year for the start of the agricultural period
agric_text = f"Cumuls de pluie enregistrés Validité : du 01 Septembre {AA} au {end_date}"
logger.info(f"Generated date line for 24h precip: '{date_line_text}'")
logger.info(f"Generated date line for cumulative precip: '{agric_text}'")

# Find the tables in the Word document
# Check if there are enough tables (at least 4: 2 for titles, 2 for data)
if len(document.tables) < 4:
    logger.critical(f"Word document '{word_template_path}' does not contain the expected number of tables (found {len(document.tables)}, expected at least 4). Exiting script.")
    sys.exit(1)

table_date_line = document.tables[0]
table_precip = document.tables[1]
table_agric_date_line = document.tables[2]
table_agri = document.tables[3]

# Insert the dynamic date line into the first cell of the first table
try:
    cell = table_date_line.cell(0, 0)
    cell.text = date_line_text
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x99, 0xDA)
    run.font.bold = True
    logger.info("Updated 24h precipitation date line in Word document.")
except Exception as e:
    logger.error(f"Error updating 24h precipitation date line in Word document (table 0, cell 0): {e}", exc_info=True)

# Insert the dynamic date line into the first cell of the third table (agric_tab)
try:
    cell_agr = table_agric_date_line.cell(0, 0)
    cell_agr.text = agric_text
    paragraph_agr = cell_agr.paragraphs[0]
    paragraph_agr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_agr = paragraph_agr.runs[0]
    run_agr.font.size = Pt(14)
    run_agr.font.color.rgb = RGBColor(0x00, 0x99, 0xDA)
    run_agr.font.bold = True
    logger.info("Updated cumulative precipitation date line in Word document.")
except Exception as e:
    logger.error(f"Error updating cumulative precipitation date line in Word document (table 2, cell 0): {e}", exc_info=True)

# 24-Hour Precipitation Table (table[1]) in Word
logger.info("Populating 24-hour precipitation table (table 1) in Word document.")
for row_index, excel_row in enumerate(ws_24h_word.iter_rows(min_row=2, values_only=True), start=1):
    if row_index >= len(table_precip.rows):
        logger.warning(f"Skipping row {row_index} for 24h precip table: Word table has fewer rows ({len(table_precip.rows)}) than Excel data.")
        break # Exit loop if Word table has no more rows
    for col, excel_col_index in enumerate([1, 3, 5]): # Columns B, D, F in Excel
        try:
            value = format_value(excel_row[excel_col_index])
            cell = table_precip.cell(row_index, col * 2 + 1) # Target columns B, D, F in Word table
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except IndexError:
            logger.error(f"Index error while processing 24h precip table at Word row {row_index}, Excel column index {excel_col_index}. Word table might be malformed. Skipping cell.")
        except Exception as e:
            logger.error(f"Error populating 24h precip table cell at Word row {row_index}, Word col {col * 2 + 1}: {e}", exc_info=True)

# Set the font size for the rows in table_precip to 9
try:
    for row in table_precip.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    logger.info("Set font size for 24h precipitation table in Word.")
except Exception as e:
    logger.error(f"Error setting font size for 24h precipitation table: {e}", exc_info=True)

# Cumulative Precipitation Table (table[3]) in Word
logger.info("Populating cumulative precipitation table (table 3) in Word document.")
for row_index, excel_row in enumerate(ws_agri_word.iter_rows(min_row=2, values_only=True), start=1):
    if row_index >= len(table_agri.rows):
        logger.warning(f"Skipping row {row_index} for cumulative precip table: Word table has fewer rows ({len(table_agri.rows)}) than Excel data.")
        break
    for col, excel_col_index in enumerate([1, 3, 5]): # Columns B, D, F in Excel
        try:
            value = format_value(excel_row[excel_col_index])
            cell = table_agri.cell(row_index, col * 2 + 1)
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except IndexError:
            logger.error(f"Index error while processing cumulative precip table at Word row {row_index}, Excel column index {excel_col_index}. Word table might be malformed. Skipping cell.")
        except Exception as e:
            logger.error(f"Error populating cumulative precip table cell at Word row {row_index}, Word col {col * 2 + 1}: {e}", exc_info=True)

# Set the font size for the rows in table_agri to 9
try:
    for row in table_agri.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    logger.info("Set font size for cumulative precipitation table in Word.")
except Exception as e:
    logger.error(f"Error setting font size for cumulative precipitation table: {e}", exc_info=True)

# Save the final Word document
try:
    document.save(output_word_path)
    logger.info(f"Combined Word document with 24-hour and cumulative precipitation saved as {output_word_path}")
except Exception as e:
    logger.error(f"Error saving final Word document to {output_word_path}: {e}", exc_info=True)

logger.info(f"Script {os.path.basename(__file__)} finished.")
