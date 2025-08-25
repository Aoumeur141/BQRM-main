import json
import locale
import os
import sys
import argparse
import subprocess

import pandas as pd
# REMOVED: from datetime import datetime, timedelta # Replaced by pytaps.date_time_utils
from openpyxl import load_workbook # Still used for wb, ws if pytaps.load_excel_workbook returns workbook object
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from docx.shared import RGBColor
import logging

# IMPORTED: pytap functions
from pytaps.data_utils import load_excel_workbook
from pytaps.file_operations import clean_directory
from pytaps.system_utils import execute_command
from pytaps.logging_utils import setup_logger
# NEW: Import only the specific date utility function needed
from pytaps.date_time_utils import get_date_n_days_ago_or_future


# --- Dynamically Determine Paths ---
SCRIPT_DIR = Path(os.path.dirname(os.path.abspath(__file__)))
PROJECT_ROOT = SCRIPT_DIR.parent

# --- Setup Logging ---
script_name = os.path.basename(__file__)

parser = argparse.ArgumentParser(description="7-create_word.py script for processing BUFR observations.")
parser.add_argument('--shared-log-file', type=str,
                    help='Path to a shared log file for chained scripts. This script expects it to be provided.')
args = parser.parse_args()

if not args.shared_log_file:
    print(f"ERROR: Script '{script_name}' received no --shared-log-file argument. "
          "This script must be run as part of a chained workflow.", file=sys.stderr)
    sys.exit(1)

logger, current_log_file_path = setup_logger(
    script_name=script_name,
    log_directory_base=SCRIPT_DIR,
    log_level=logging.INFO,
    shared_log_file_path=args.shared_log_file
)

logger.info(f"--- Script '{script_name}' execution started. ---")
logger.info(f"Logger configured. Logs will be saved to: {current_log_file_path}")
logger.info(f"Current script directory: {SCRIPT_DIR}")
logger.info(f"Determined project root directory: {PROJECT_ROOT}")
logger.info(f"Current working directory (before change): {os.getcwd()}")

try:
    os.chdir(PROJECT_ROOT)
    logger.info(f"Changed current working directory to: {os.getcwd()}")
except OSError as e:
    logger.critical(f"Failed to change directory to {PROJECT_ROOT}: {e}")
    logger.critical("Exiting script due to critical directory change failure.")
    sys.exit(1)

# --- Date Preparation using pytaps.date_time_utils ---
try:
    # Get today's date formatted as 'DD-MM-YYYY'
    date_today = get_date_n_days_ago_or_future(n_days=0, format_string='%d-%m-%Y', logger_instance=logger)

    # Get tomorrow's date formatted as 'DD-MM-YYYY'
    date_demain = get_date_n_days_ago_or_future(n_days=1, format_string='%d-%m-%Y', logger_instance=logger)

    # Get yesterday's date formatted as 'DD-MM-YYYY'
    date_yesterday = get_date_n_days_ago_or_future(n_days=-1, format_string='%d-%m-%Y', logger_instance=logger)

    logger.info(f"Processed dates - Today: {date_today}, Tomorrow: {date_demain}, Yesterday: {date_yesterday}")

except Exception as e:
    logger.critical(f"Error during date preparation using pytaps.date_time_utils.get_date_n_days_ago_or_future: {e}")
    logger.exception("Full traceback for date preparation error:")
    logger.critical("Exiting script due to date preparation failure.")
    sys.exit(1)


# --- Clean up temporary files ---
tmp_dir = PROJECT_ROOT / "tmp"
logger.info(f"Attempting to clean up temporary files in: {tmp_dir} using pytaps.file_operations.clean_directory.")
try:
    clean_directory(tmp_dir, logger_instance=logger)
    logger.info("Temporary files cleanup complete.")
except Exception as e:
    logger.error(f"Temporary file cleanup failed: {e}. Continuing script execution.")
    logger.exception("Full traceback for temporary file cleanup error:")


# Fonction pour modifier du texte dans un paragraphe tout en gardant son style
def modify_paragraph(doc, search_text, new_text):
    """
    Modifies text within a paragraph in a Word document while preserving its style.
    """
    logger.info(f"Attempting to modify paragraph: search_text='{search_text}', new_text='{new_text}'")
    found = False
    for para in doc.paragraphs:
        if search_text in para.text:
            found = True
            if para.runs:
                run = para.runs[0]
                font = run.font

                # Sauvegarde des styles existants
                font_name = font.name
                font_size = font.size
                font_bold = font.bold
                font_italic = font.italic
                font_color = font.color.rgb if font.color else None

                # Effacer et réécrire le paragraphe avec les mêmes styles
                para.clear()
                new_run = para.add_run(new_text)

                new_run.font.name = font_name
                new_run.font.size = font_size
                new_run.font.bold = font_bold
                new_run.font.italic = font_italic
                if font_color:
                    new_run.font.color.rgb = font_color
                logger.info(f"Paragraph successfully modified: '{search_text}' replaced with '{new_text}'")
            else:
                # If there are no runs, just replace the text (style might be default)
                para.clear()
                para.add_run(new_text)
                logger.warning(f"Paragraph '{search_text}' found but had no runs. Text replaced, but style might be default.")
            break
    if not found:
        logger.warning(f"Search text '{search_text}' not found in any paragraph.")


# Chemins des fichiers
xlsx_path = PROJECT_ROOT / "outputs/tab_reg/all_stations.xlsx"
word_template_path = PROJECT_ROOT / "template/Bulletin_sud.docx"
output_word_path = PROJECT_ROOT / f"bulletins/Bulletin_antiacridienne_{date_today}.docx" # Using date_today for filename

logger.info(f"Excel file path: {xlsx_path}")
logger.info(f"Word template path: {word_template_path}")
logger.info(f"Output Word document path: {output_word_path}")

# Chargement des fichiers Excel
try:
    # Using pytaps.data_utils.load_excel_workbook to load the workbook object
    loaded_data = load_excel_workbook(xlsx_path, logger_instance=logger)
    wb = loaded_data[0] # Assuming the workbook object is the first element of the returned tuple
    ws = wb.active
    logger.info(f"Successfully loaded Excel workbook using pytaps: {xlsx_path}")
except FileNotFoundError:
    logger.critical(f"Excel file not found at: {xlsx_path}")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading Excel workbook {xlsx_path}: {e}")
    logger.exception("Full traceback for Excel workbook loading error:")
    sys.exit(1)

# Load the Excel file containing station data into a DataFrame (still using pandas for this)
try:
    df = pd.read_excel(xlsx_path)
    logger.info(f"Successfully loaded data into Pandas DataFrame from: {xlsx_path}")
    if df.empty:
        logger.warning("DataFrame loaded from Excel is empty.")
except Exception as e:
    logger.critical(f"Error reading Excel into DataFrame {xlsx_path}: {e}")
    logger.exception("Full traceback for DataFrame loading error:")
    sys.exit(1)

# Chargement du modèle Word
try:
    document = Document(word_template_path)
    logger.info(f"Successfully loaded Word template: {word_template_path}")
except FileNotFoundError:
    logger.critical(f"Word template file not found at: {word_template_path}")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading Word template {word_template_path}: {e}")
    logger.exception("Full traceback for Word template loading error:")
    sys.exit(1)


# Mise à jour des dates dans le document Word
modify_paragraph(document,
                 "Tableau de Températures observées pour le YESTERDAY et prévues pour le TODAY et DEMAIN :",
                 f"Tableau de Températures observées pour le {date_yesterday} et prévues pour le {date_today} et {date_demain} : ")

# Function to update the existing table in the Word document with data from the Excel file
def update_existing_table(document, df):
    """
    Updates the second table in the Word document with data from the DataFrame.
    Applies consistent font size, cell width, and centers content.
    """
    logger.info("Starting to update the Word document table.")
    try:
        # Find the table in the document (assuming it's the second table, index 1)
        if len(document.tables) < 2:
            logger.error("Less than 2 tables found in the document. Cannot update the target table.")
            return False
        table = document.tables[1]
        logger.info(f"Found table for update (index 1). Table has {len(table.rows)} rows and {len(table.columns)} columns initially.")

        # Ensure the table has enough rows to accommodate the data
        # Excluding the first 2 rows (merged header)
        required_rows = len(df) + 2 # data rows + 2 header rows
        current_rows = len(table.rows)
        while current_rows < required_rows:
            table.add_row()
            current_rows += 1
            logger.debug(f"Added a new row to the table. Current rows: {current_rows}")

        logger.info(f"Table now has {current_rows} rows, sufficient for {len(df)} data entries.")

        # Iterate over rows of the dataframe and fill the table with data
        for index, row in df.iterrows():
            # Find the row in the table to populate (start from the 3rd row, index 2)
            row_cells = table.rows[index + 2].cells # +2 to skip the first two rows (merged header)

            # Populate each cell with the corresponding data from the dataframe
            # Skip the first column (Station) from df, start from col_index 1 in Word table
            for col_index, col_name in enumerate(df.columns[1:], start=1):
                value = row[col_name]
                if col_index < len(row_cells): # Ensure we don't exceed the available columns in the table
                    cell_text = str(value) if pd.notna(value) else ""
                    row_cells[col_index].text = cell_text
                    row_cells[col_index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logger.debug(f"Row {index+2}, Col {col_index}: Set text to '{cell_text}' and centered.")
                else:
                    logger.warning(f"Skipping column {col_name} for row {index} as table does not have enough columns at index {col_index}.")

        # Apply consistent font size, cell width, and row height
        logger.info("Applying consistent formatting to table cells.")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                # Adjust font size for all cells
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8) # Font size set to 8 to fit the table properly
                        logger.debug(f"Row {r_idx}, Col {c_idx}: Set font size to 8pt.")

                # Set consistent column widths
                if c_idx == 0: # First column (Station)
                    cell.width = Inches(1.5)
                else: # Other columns
                    cell.width = Inches(0.5)
                logger.debug(f"Row {r_idx}, Col {c_idx}: Set width to {cell.width.inches} inches.")

                # Center the content inside the cell (already done above for data cells, but ensure for all)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_after = Pt(0) # Remove extra space after paragraphs
                    paragraph.paragraph_format.space_before = Pt(0) # Remove extra space before paragraphs
                    logger.debug(f"Row {r_idx}, Col {c_idx}: Centered text and removed paragraph spacing.")
        logger.info("Table update and formatting complete.")
        return True
    except IndexError:
        logger.error("Table index out of range. Check if the table exists at the expected index (1).")
        return False
    except Exception as e:
        logger.error(f"An error occurred during table update: {e}", exc_info=True)
        return False

# Update the existing table with data from the Excel file
if not update_existing_table(document, df):
    logger.critical("Failed to update the Word document table. Exiting.")
    sys.exit(1)

# Save the updated Word document
try:
    document.save(output_word_path)
    logger.info(f"Word document successfully saved as: {output_word_path}")
except Exception as e:
    logger.critical(f"Failed to save Word document to {output_word_path}: {e}", exc_info=True)
    sys.exit(1)

# --- Run Next Program ---
logger.info(f"Attempting to run next program: 8-send_MSG.py")
next_script_path = SCRIPT_DIR / "8-send_MSG.py"
try:
    execute_command(
        [sys.executable, str(next_script_path), "--shared-log-file", current_log_file_path],
        cwd=SCRIPT_DIR, # Ensure the command is executed from the SCRIPT_DIR
        logger_instance=logger
    )
    logger.info("Successfully executed 8-send_MSG.py.")
except (subprocess.CalledProcessError, FileNotFoundError) as e:
    logger.exception(f"An error occurred while running 8-send_MSG.py: {e}")
    sys.exit(1)
except Exception as e:
    logger.exception(f"An unexpected error occurred in the main script flow: {e}")
    sys.exit(1)

logger.info(f"--- Script '{script_name}' finished. ---")
