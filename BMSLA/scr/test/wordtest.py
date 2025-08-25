import os
import pandas as pd
from datetime import datetime, timedelta
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
#from modules.module import run_next_program


source_dir = "/home/ibousri/AROME_DUST/Projet_Nationale/Sonelgaz/sonelgaz_oper/"

# Function to retrieve today's and tomorrow's date
def prepare_date():
    #today = datetime.now()
    today = datetime(2025, 4, 10)  # Example of a fixed date for testing
    tomorrow = today + timedelta(days=1)

    date_today = today.strftime('%d-%m-%Y')   
    date_tomorrow = tomorrow.strftime('%d-%m-%Y')   
    
    return date_today, date_tomorrow 

# Get today's and tomorrow's dates
date_today, date_tomorrow = prepare_date()

# Function to modify text in a paragraph while preserving its style
def modify_paragraph(doc, search_text, new_text):
    for para in doc.paragraphs:
        if search_text in para.text:
            if para.runs:
                run = para.runs[0]
                font = run.font
                
                # Save existing styles
                font_name = font.name
                font_size = font.size
                font_bold = font.bold
                font_italic = font.italic
                font_color = font.color.rgb if font.color else None

                # Clear and rewrite the paragraph with the same styles
                para.clear()
                new_run = para.add_run(new_text)

                new_run.font.name = font_name
                new_run.font.size = font_size
                new_run.font.bold = font_bold
                new_run.font.italic = font_italic
                if font_color:
                    new_run.font.color.rgb = font_color

            break  

# File paths (these can be modified as per your requirement)
xlsx_path = os.path.join(source_dir, "Arpg/all_stations.xlsx")  # The path to the generated Excel file
word_template_path = os.path.join(source_dir, "template/Bulletin_sud.docx")
output_word_path = os.path.join(source_dir, f"bulletins/Bulletin_Sud_{date_today}.docx")

# Load the Excel file containing station data (from the previous script)
df = pd.read_excel(xlsx_path)

# Load the Word template
document = Document(word_template_path)

# Update the dates in the Word document
modify_paragraph(document, "Bulletin du  DATE_TODAY", f"Bulletin du {date_today}")
#modify_paragraph(document, "Prévision valable pour la journée du : DATE_DEMAIN", f"Prévision valable pour la journée du : {date_tomorrow}")

# Function to update the existing table in the Word document with data from the Excel file
def update_existing_table(document, df):
    # Find the table in the document (assuming it's the first table)
    table = document.tables[1]  # Adjust if it's not the first table
    
    # Ensure the table has enough rows to accommodate the data
    while len(table.rows) - 2 < len(df):  # Excluding the first 3 rows (merged header)
        table.add_row()

    # Iterate over rows of the dataframe and fill the table with data
    for index, row in df.iterrows():
        # Find the row in the table to populate (start from the 4th row, index 3)
        row_cells = table.rows[index + 2].cells  # +3 to skip the first three rows (merged header)

        # Populate each cell with the corresponding data from the dataframe
        for col_index, col_name in enumerate(df.columns[1:], start=1):  # Skip the first column (Station)
            value = row[col_name]
            if col_index < len(row_cells):  # Ensure we don't exceed the available columns in the table
                row_cells[col_index].text = str(value) if pd.notna(value) else ""
                row_cells[col_index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Apply consistent font size, cell width, and row height
    for row in table.rows:
        for col_index, cell in enumerate(row.cells):
            # Adjust font size for all cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)  # Font size set to 8 to fit the table properly

            # Set consistent column widths
            if col_index == 0:  # First column (Station)
                cell.width = Inches(1.5)
            else:  # Other columns
                cell.width = Inches(0.5)
            
            # Adjust row height to avoid excessive space
            #cell._element.get_or_add_tr().set('height', '200')  # Adjust row height in points

            # Center the content inside the cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align text

    # Reduce space between rows (if needed)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)  # Remove extra space after paragraphs
                paragraph.paragraph_format.space_before = Pt(0)  # Remove extra space before paragraphs

# Update the existing table with data from the Excel file
update_existing_table(document, df)

# Save the updated Word document
document.save(output_word_path)
print(f"Word document saved as {output_word_path}")

# Run the next program (as specified in the original script)
#os.chdir(os.path.join(source_dir, "scr"))
#run_next_program("8-send_MSG.py")
