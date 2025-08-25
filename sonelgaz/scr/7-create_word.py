import os
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from docx.shared import RGBColor # Keep this for modify_paragraph if it uses it directly

import logging
import argparse
import sys
import subprocess

# IMPORTED: pytap functions
from pytaps.data_utils import load_excel_workbook
from pytaps.file_operations import clean_directory
from pytaps.system_utils import execute_command
from pytaps.logging_utils import setup_logger
from pytaps.date_time_utils import get_date_n_days_ago_or_future

# --- Dynamically Determine Paths ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# Define the project root, assuming SCRIPT_DIR (e.g., 'scr') is a subdirectory of the main project folder.
PROJECT_ROOT = Path(SCRIPT_DIR).parent

# --- Setup Logging ---
script_name = os.path.basename(__file__)

# Parse command-line arguments for a shared log file path
parser = argparse.ArgumentParser(description="7-create_word.py script for processing BUFR observations.")
parser.add_argument('shared_path_log_file', type=str,
                    help='Path to a shared log file for chained scripts. This script expects it to be provided.')
args = parser.parse_args()

# Configure the logger. This script expects a shared log file path.
if not args.shared_path_log_file:
    # Print to stderr because the logger might not be fully set up yet
    print(f"ERROR: Script '{script_name}' received no shared_path_log_file argument. "
          "This script must be run as part of a chained workflow.", file=sys.stderr)
    sys.exit(1)

logger, current_log_file_path = setup_logger(
    script_name=script_name,
    log_directory_base=SCRIPT_DIR,
    log_level=logging.INFO,
    shared_log_file_path=args.shared_path_log_file
)

logger.info(f"--- Script '{script_name}' execution started. ---")
logger.info(f"Logger configured. Logs will be saved to: {current_log_file_path}")
logger.info(f"Current script directory: {SCRIPT_DIR}")
logger.info(f"Determined project root directory: {PROJECT_ROOT}")
logger.info(f"Current working directory (before change): {os.getcwd()}")

try:
    os.chdir(PROJECT_ROOT)
    logger.info(f"Changed current working directory to: {os.getcwd()}")
except OSError as e:
    logger.critical(f"Failed to change directory to {PROJECT_ROOT}: {e}")
    logger.critical("Exiting script due to critical directory change failure.")
    sys.exit(1)

# --- Date Preparation ---
# Replaced the call to the old function with pytaps.date_utils.get_date_n_days_ago_or_future
try:
   # Get today's date formatted as 'DD-MM-YYYY'
   date_today = get_date_n_days_ago_or_future(n_days=0, format_string='%d-%m-%Y', logger_instance=logger)

   # Get tomorrow's date formatted as 'DD-MM-YYYY'
   date_demain = get_date_n_days_ago_or_future(n_days=1, format_string='%d-%m-%Y', logger_instance=logger)

   logger.info(f"Processed dates - Today: {date_today}, Tomorrow: {date_demain}")

except Exception as e:
   logger.critical(f"Error during date preparation using pytaps.date_utils.get_date_n_days_ago_or_future: {e}")
   logger.exception("Full traceback for date preparation error:")
   logger.critical("Exiting script due to date preparation failure.")
   sys.exit(1)



# --- Word Document Helper Functions (Moved back into this script) ---

def modify_paragraph(doc, search_text, new_text):
    """
    Modifies text within a paragraph in a Word document while preserving its style.
    This function is directly from the old version, enhanced with logging.
    """
    logger.debug(f"Searching for paragraph containing '{search_text}' to replace with '{new_text}'.")
    found = False
    for para_idx, para in enumerate(doc.paragraphs):
     ## so we have llop that get stored the text and there index from doc so using the modules paragraph from libreries docs to get the paragraph and then use it in enumerate to get it and here index as output.
        if search_text in para.text:
            
            logger.info(f"Found paragraph containing '{search_text}' at index {para_idx}. Attempting to modify.")
            if para.runs:
                run = para.runs[0]
                font = run.font

                # Sauvegarde des styles existants
                font_name = font.name
                font_size = font.size
                font_bold = font.bold
                font_italic = font.italic
                font_color = font.color.rgb if font.color else None
                logger.debug(f"Preserving font styles: Name='{font_name}', Size={font_size}, Bold={font_bold}, Italic={font_italic}, Color={font_color}.")

                # Effacer et réécrire le paragraphe avec les mêmes styles
                para.clear()
                new_run = para.add_run(new_text)

                new_run.font.name = font_name
                new_run.font.size = font_size
                new_run.font.bold = font_bold
                new_run.font.italic = font_italic
                if font_color:
                    new_run.font.color.rgb = font_color
                logger.info(f"Successfully modified paragraph from '{search_text}' to '{new_text}'.")
                found = True
            else:
                logger.warning(f"Paragraph containing '{search_text}' at index {para_idx} found but has no runs to modify. Clearing and adding new run without preserving style.")
                para.clear()
                para.add_run(new_text)
                found = True
            break # Stop after the first match
    if not found:
        logger.warning(f"Search text '{search_text}' not found in any paragraph. No modification performed.")

def update_table(table, worksheet, table_name="Unnamed Table"):
    """
    Updates a Word table with data from an Excel worksheet.
    Applies specific formatting (font size, column width, alignment).
    This function is directly from the old version, enhanced with logging.
    """
    logger.info(f"Starting update for Word table: '{table_name}' from Excel worksheet: '{worksheet.title}'.")

    try:
        table.style.paragraph_format.space_after = 0
        logger.debug(f"Set paragraph space_after to 0 for table '{table_name}'.")
    except Exception as e:
        logger.warning(f"Could not set paragraph_format.space_after for table '{table_name}': {e}")
        logger.debug(f"Traceback for paragraph_format warning for table '{table_name}':", exc_info=True)


    # Iterate Excel rows, starting from the second row (min_row=2) as per old code
    for row_index, excel_row_values in enumerate(worksheet.iter_rows(min_row=2, values_only=True), start=2):
        logger.debug(f"Processing Excel row {row_index} for table '{table_name}'.")
        for col_offset, excel_col_index in enumerate(range(1, 15)): # Excel columns B to O (1 to 14)
            # The old code used `table.cell(row_index, col+1)`.
            # `row_index` starts at 2. `col+1` starts at 1.
            # This means Excel row 2 -> Word table row index 2.
            # Excel col 1 (B) -> Word table col index 2.
            # This implies the Word table has a header at row 0, and data starts from row 1.
            # The old code's logic `table.cell(row_index, col+1)` directly maps
            # Excel's 1-based row index (from `start=2`) and Excel's 1-based column index
            # (from `range(1,15)`) to the Word table's 0-based indices.
            # Let's ensure bounds checking.
            word_table_col_index = col_offset + 1 # This maps Excel's 1-based column (B) to Word's 0-based index (1)

            if row_index < len(table.rows): # Check if the target row exists in Word table
                if word_table_col_index < len(table.rows[row_index].cells): # Check if the target column exists in Word table
                    cell = table.cell(row_index, word_table_col_index)
                    cell_value = excel_row_values[excel_col_index] # Get value from Excel row
                    cell.text = str(cell_value) if cell_value is not None else ""
                    logger.debug(f"  Cell ({row_index},{word_table_col_index}) updated with: '{cell.text}' (from Excel col {excel_col_index}).")
                else:
                    logger.warning(f"Skipping cell update for Word table '{table_name}' at row {row_index}, col {word_table_col_index}: Cell out of bounds. (Excel col {excel_col_index})")
            else:
                logger.warning(f"Skipping update for Word table '{table_name}': No corresponding Word table row for Excel row {row_index}.")
                break # No more rows in Word table to update for this worksheet

    # Apply width, font size and alignment
    logger.info(f"Applying formatting to table: '{table_name}'.")
    for row_idx, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)  # Font size 9 for all columns
            logger.debug(f"  Cell ({row_idx},{col_index}) font size set to 9pt.")

            if col_index == 0:
                cell.width = Inches(1.5)  # Width for "Station" column
                logger.debug(f"  Cell ({row_idx},{col_index}) width set to 1.5 inches (Station column).")
            else:
                cell.width = Inches(0.5)  # Width for other columns
                logger.debug(f"  Cell ({row_idx},{col_index}) width set to 0.5 inches.")

            # Ensure paragraph exists before setting alignment
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centered alignment
                logger.debug(f"  Cell ({row_idx},{col_index}) alignment set to center.")
            else:
                logger.warning(f"Cell at row {row_idx}, col {col_index} in '{table_name}' has no paragraphs to align.")
    logger.info(f"Finished updating and formatting table: '{table_name}'.")


# Chemins des fichiers (Corrected to use PROJECT_ROOT for outputs and template)
logger.info("Defining file paths for Excel, Word template, and output.")
xlsx_Ouest_path = PROJECT_ROOT / "outputs" / "tab_reg" / "regions_Ouest.xlsx"
xlsx_Centre_path = PROJECT_ROOT / "outputs" / "tab_reg" / "regions_Centre.xlsx"
xlsx_Est_path = PROJECT_ROOT / "outputs" / "tab_reg" / "regions_Est.xlsx"
xlsx_Sud_path = PROJECT_ROOT / "outputs" / "tab_reg" / "regions_Sud.xlsx"
word_template_path = PROJECT_ROOT / "template" / "Bulletin.docx"
output_word_path = PROJECT_ROOT / "bulletins" / f"Bulletin_Sonelgaz_{date_today}.docx"

logger.info(f"Excel Ouest path: {xlsx_Ouest_path}")
logger.info(f"Excel Centre path: {xlsx_Centre_path}")
logger.info(f"Excel Est path: {xlsx_Est_path}")
logger.info(f"Excel Sud path: {xlsx_Sud_path}")
logger.info(f"Word template path: {word_template_path}")
logger.info(f"Output Word document path: {output_word_path}")


# Chargement des fichiers Excel
logger.info("Loading Excel workbooks using pytaps.data_utils.load_excel_workbook...")
try:
    wb_ouest, ws_ouest = load_excel_workbook(xlsx_Ouest_path, logger_instance=logger)
    wb_centre, ws_centre = load_excel_workbook(xlsx_Centre_path, logger_instance=logger)
    wb_est, ws_est = load_excel_workbook(xlsx_Est_path, logger_instance=logger)
    wb_sud, ws_sud = load_excel_workbook(xlsx_Sud_path, logger_instance=logger)
    logger.info("All Excel workbooks loaded successfully.")
except Exception as e:
    logger.critical(f"Failed to load one or more Excel workbooks: {e}. Exiting script.")
    logger.exception("Full traceback for Excel loading failure:")
    sys.exit(1)


# Chargement du modèle Word
logger.info("Loading Word template...")
document = None
try:
    document = Document(word_template_path)
    logger.info(f"Successfully loaded Word template: {word_template_path}")
except FileNotFoundError:
    logger.critical(f"Error: Word template not found at {word_template_path}.")
    sys.exit(1)
except Exception as e:
    logger.critical(f"Error loading Word template {word_template_path}: {e}")
    logger.exception("Full traceback for Word template loading error:")
    sys.exit(1)


# Mise à jour des dates dans le document Word
logger.info("Updating dates in the Word document using local modify_paragraph function.")
try:
    modify_paragraph(document, "Bulletin du  DATE_TODAY", f"Bulletin du {date_today}")
    modify_paragraph(document, "Prévision valable pour la journée du : DATE_DEMAIN", f"Prévision valable pour la journée du : {date_demain}")
    logger.info("Dates in Word document updated successfully.")
except Exception as e:
    logger.error(f"Error updating dates in Word document: {e}. Continuing with script, but output might be incorrect.")
    logger.exception("Full traceback for Word date modification error:")


# Mise à jour des tableaux dans le document Word
logger.info("Updating all tables in the Word document with Excel data using local update_table function.")
if document and len(document.tables) >= 9: # Check if enough tables exist
    try:
        update_table(document.tables[2], ws_ouest, "Ouest Region Table")
        update_table(document.tables[4], ws_centre, "Centre Region Table")
        update_table(document.tables[6], ws_est, "Est Region Table")
        update_table(document.tables[8], ws_sud, "Sud Region Table")
        logger.info("All specified tables in Word document updated successfully.")
    except Exception as e:
        logger.critical(f"Error updating one or more tables in Word document: {e}. Exiting script.")
        logger.exception("Full traceback for Word table update error:")
        sys.exit(1)
else:
    logger.critical(f"Not enough tables found in the Word document. Expected at least 9, found {len(document.tables) if document else 'N/A'}.")
    sys.exit(1)


# Sauvegarde du document mis à jour
logger.info(f"Attempting to save the updated Word document to: {output_word_path}")
try:
    output_word_path.parent.mkdir(parents=True, exist_ok=True)
    logger.debug(f"Ensured output directory exists: {output_word_path.parent}")
    document.save(output_word_path)
    logger.info(f"Word document successfully saved as {output_word_path}")

    # --- ADDED CONFIRMATION CHECK ---
    if output_word_path.exists():
        logger.info(f"CONFIRMATION: File '{output_word_path.name}' exists on disk at '{output_word_path.parent}' (size: {output_word_path.stat().st_size} bytes).")
    else:
        logger.critical(f"CRITICAL ERROR: File was reported saved but does NOT exist on disk at {output_word_path}. This is highly unusual.")
        sys.exit(1)
    # --- END ADDED CONFIRMATION CHECK ---

except Exception as e:
    logger.critical(f"Error saving Word document to {output_word_path}: {e}")
    logger.exception("Full traceback for Word document saving error:")
    sys.exit(1)

# --- Clean up temporary files ---
tmp_dir = PROJECT_ROOT / "tmp" # Use PROJECT_ROOT for tmp directory
logger.info(f"Attempting to clean up temporary files in: {tmp_dir} using pytap.file_operations.clean_directory.")
try:
    #clean_directory(tmp_dir, logger_instance=logger)
    logger.info("Temporary files cleanup complete.")
except Exception as e:
    logger.error(f"Temporary file cleanup failed: {e}. Continuing script execution.")
    logger.exception("Full traceback for temporary file cleanup error:")


# --- Run Next Program ---
logger.info(f"Attempting to run next program: 8-send_MSG.py")
next_script_path = os.path.join(SCRIPT_DIR, "8-send_MSG.py") 
try:
    execute_command(
        [sys.executable, next_script_path, current_log_file_path],
        cwd=SCRIPT_DIR
    )
    logger.info("Successfully executed 8-send_MSG.py.")
except (subprocess.CalledProcessError, FileNotFoundError) as e:
    logger.exception(f"An error occurred while running 8-send_MSG.py: {e}")
    sys.exit(1)
except Exception as e:
    logger.exception(f"An unexpected error occurred in the main script flow: {e}")
    sys.exit(1)

logger.info(f"--- Script '{script_name}' finished. ---")
